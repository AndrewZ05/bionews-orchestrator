#!/usr/bin/env python3
"""
Build exports/site_and_condition_model.docx from docs/SITE_AND_CONDITION_MODEL.md.

The Markdown is the source of truth -- it is version-controlled and its queries
are dry-run validated against BigQuery. This renders the same content as a Word
document for a senior-management review, so the two cannot drift: edit the
Markdown, re-run this.

  python scripts/build_site_condition_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
SOURCE = REPO / "docs" / "SITE_AND_CONDITION_MODEL.md"
EXPORTS = REPO / "exports"
OUT = EXPORTS / "site_and_condition_model.docx"

NAVY = RGBColor(0x0D, 0x47, 0xA1)
GREY = RGBColor(0x5F, 0x63, 0x68)
CODE_BG = RGBColor(0x1F, 0x2A, 0x37)
RED = RGBColor(0xEE, 0x00, 0x00)

# Passages the reader marked in red in a previous review. Regenerating the
# document would otherwise silently discard those marks, so they are carried in
# the Markdown as an explicit <<red>> ... <</red>> span and re-applied here.
RED_OPEN = "<<red>>"
RED_CLOSE = "<</red>>"


def _clean(text: str) -> str:
    """Strip inline markdown. A printed heading should not show backticks."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)


def _style_heading(h, color=NAVY):
    for r in h.runs:
        r.font.color.rgb = color
    return h


def _rich(paragraph, text, size=11, red=False):
    """Render inline **bold**, `code` and [link](x) as plain runs.

    Markdown emphasis markers are stripped rather than carried through -- a
    printed document does not need backticks, and leaving them in reads as
    typos to anyone outside the data team.
    """
    for part in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(size - 1)
        else:
            # Collapse [text](link) to just the text.
            r = paragraph.add_run(re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", part))
        r.font.size = Pt(size)
        if red:
            r.font.color.rgb = RED


def _para(doc, text, size=11, after=8, red=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    _rich(p, text, size, red=red)
    return p


def _bullet(doc, text, red=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    _rich(p, text, red=red)
    return p


def _code(doc, lines):
    """SQL block: monospace, tight leading, on a shaded background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent = Pt(14)
    for i, line in enumerate(lines):
        r = p.add_run(line + ("\n" if i < len(lines) - 1 else ""))
        r.font.name = "Consolas"
        r.font.size = Pt(8.5)
        r.font.color.rgb = CODE_BG
    return p


def _table(doc, header, rows):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Light List Accent 1"
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
        if i > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            # Strip inline markdown inside table cells too.
            v = re.sub(r"\*\*(.+?)\*\*", r"\1", str(v))
            v = re.sub(r"`([^`]+)`", r"\1", v)
            r = cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(9.5)
            if i > 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def _split_table_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def build():
    md = SOURCE.read_text(encoding="utf-8").split("\n")
    doc = Document()

    i = 0
    in_code = False
    in_red = False
    code_lines: list[str] = []

    while i < len(md):
        line = md[i]

        # ---- reader's red markers (whole-block, not inline) ----------------
        if line.strip() == RED_OPEN:
            in_red = True
            i += 1
            continue
        if line.strip() == RED_CLOSE:
            in_red = False
            i += 1
            continue

        # ---- fenced code -------------------------------------------------
        if line.startswith("```"):
            if in_code:
                _code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ---- tables ------------------------------------------------------
        if (
            line.startswith("|")
            and i + 1 < len(md)
            and set(md[i + 1].replace("|", "").strip()) <= set("-: ")
        ):
            header = _split_table_row(line)
            i += 2
            rows = []
            while i < len(md) and md[i].startswith("|"):
                rows.append(_split_table_row(md[i]))
                i += 1
            _table(doc, header, rows)
            continue

        # ---- headings ----------------------------------------------------
        if line.startswith("# "):
            _style_heading(doc.add_heading(_clean(line[2:].strip()), level=0))
        elif line.startswith("## "):
            _style_heading(
                doc.add_heading(_clean(re.sub(r"^\d+\.\s*", "", line[3:].strip())), level=1)
            )
        elif line.startswith("### "):
            _style_heading(
                doc.add_heading(_clean(line[4:].strip()), level=2),
                RED if in_red else NAVY,
            )
        elif line.startswith("- "):
            _bullet(doc, line[2:].strip(), red=in_red)
        elif line.strip() == "---":
            pass  # section rules are carried by the headings themselves
        elif line.strip():
            _para(doc, line.strip(), red=in_red)

        i += 1

    EXPORTS.mkdir(exist_ok=True)
    doc.save(OUT)
    print(f"[OK] {OUT.relative_to(REPO)}")
    return OUT


if __name__ == "__main__":
    build()
