"""Generate BI-team query guides as .docx for the profile database and identity hub.

Plain text only: no emojis, no shading, no color. Bold for emphasis, monospace
for SQL.

Every query in here is validated against production by
tests/unit/test_bi_query_docx.py, which dry-runs each one. A query that stops
being valid fails CI rather than reaching an analyst.

Usage:
    python scripts/build_bi_query_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

EXPORTS = Path(__file__).resolve().parents[1] / "exports"
PROFILE_OUT = EXPORTS / "profile_db_bi_queries.docx"
HUB_OUT = EXPORTS / "identity_hub_bi_queries.docx"


# ---------------------------------------------------------------------------
# Profile database: known vs unknown, conditions, roles, reachability
# ---------------------------------------------------------------------------
PROFILE_INTRO = [
    "This is the query guide for the profile database. Every query has been run "
    "against production and returns what it claims to.",
    "Start with profile_metrics. It is one row per person with every headline "
    "metric already computed as a true/false column, so you group it rather than "
    "rebuilding a definition. That matters because each metric below has a trap "
    "in it, and every one of those traps has produced a wrong number here before.",
]

# Text prefixed with RED:: renders dark red. Marks passages added since the
# reader last reviewed the document, so a re-read can go straight to what
# changed. The prefix is stripped before rendering; drop it once reviewed.
RED = "RED::"
_DARK_RED = RGBColor(0xC0, 0x00, 0x00)


PROFILE_RULES = [
    (
        "Known versus unknown is the first filter, not an afterthought",
        "The database holds about 7.5 million profiles, but only about 799,000 "
        "are people you can name. The rest are cookie-only records with no email, "
        "no login and no durable identifier -- closer to browsers than humans. "
        "is_known_person is the flag. Quoting the full 7.5 million as an audience "
        "overstates reality roughly eight times.",
    ),
    (
        RED + "The HCP number has three sizes -- pick the right one",
        RED + "is_verified_hcp counts 380,344 clinicians whose credential is "
        "real and current. It is not an audience. About 73 percent of them came "
        "from the federal NPI registry rather than from our audience -- 269,749 "
        "hold nothing but an email address and an NPI number -- and have "
        "never produced a pageview, an open or a click. Quoting it "
        "commercially overstates reachable clinicians roughly sixteen times "
        "-- the same trap as quoting 7.5 million profiles as people. Use "
        "is_verified_hcp AND is_mailable (23,019) for anything commercial or "
        "external: verified, subscribed, opted in, not consent-denied. It is "
        "the only figure that survives the question can you actually reach "
        "them. Use is_engaged_hcp (102,217) for how many clinicians do we "
        "know -- verified, and has at some point produced a pageview, a human "
        "email open or a click.",
    ),
    (
        RED + "is_engaged_hcp says ever, not lately, and it is a floor",
        RED + "There is no time bound on it. Only 47,802 of the 102,217 were "
        "active in the last 90 days, so for a live audience use "
        "is_verified_hcp AND is_active_email_90d (13,934) instead. It is also "
        "82 percent pageview-driven, and a pageview only counts once identity "
        "resolution has tied the visit to a person -- clinicians we failed to "
        "resolve are invisible to it. Treat the number as conservative: the "
        "true engaged population is somewhat higher, never lower.",
    ),
    (
        RED + "Mailable counts went up on 2026-08-20, and nobody was newly opted in",
        RED + "is_mailable, is_active_email_90d and is_active_email_30d each rose "
        "by about 5,200. No consent changed. consent_status is written from the "
        "OneTrust analytics cookie group seen in GA4 -- it records whether "
        "someone accepted an analytics cookie, not whether they agreed to be "
        "emailed -- and it was gating the email metrics. That excluded 5,220 "
        "people who had explicitly opted in and were subscribed in Mailchimp, "
        "350 of them verified HCPs. Mailchimp mailed them anyway, so the flag "
        "protected nobody and only made our numbers disagree with the sending "
        "platform. An explicit opt-in now stands. If you are comparing to a "
        "figure pulled before 2026-08-20, this is the difference.",
    ),
    (
        RED + "A lapsed credential is not a lapsed person",
        RED + "has_deactivated_npi counts 4,170 clinicians whose NPI the "
        "federal registry has retired -- they retired, changed entity type or "
        "died. They stopped counting as is_verified_hcp on 2026-08-19, but "
        "they remain in the database with their history intact and 95 percent "
        "of them still read. Exclude them from clinician counts; do not "
        "exclude them from audience work.",
    ),
    (
        "Role flags do not add up, and that is correct",
        "A person can be both a healthcare professional and a patient, so the role "
        "columns sum to more than the profile count. Around 3,500 people hold two "
        "or more roles. Never present them as a pie chart. When you need mutually "
        "exclusive buckets use primary_role, which applies the priority "
        "hcp, patient, caregiver, family, other.",
    ),
    (
        "A missing role is not a data problem",
        "About 92 percent of profiles carry no role at all, because most site "
        "visitors never tell us who they are. A persona chart that includes them "
        "is mostly 'unspecified' and says nothing. Filter to has_any_role, and "
        "state the denominator you used.",
    ),
    (
        "Filter conditions on the normalized column",
        "The raw condition label carries both short codes and full names for the "
        "same disease, so filtering on it splits an audience in half. Use "
        "condition_key or condition_label from profile_metrics, which resolve "
        "every variant to one canonical value.",
    ),
    (
        "created_at is not a signup date",
        "created_at records when the pipeline first observed a profile, not when "
        "the person arrived. Use has_registration_date and acquisition_month, "
        "which exclude cookie sightings automatically, or the "
        "profile_growth_known view.",
    ),
    (
        "communication_opt_in is not consent",
        "OneTrust access has been pending since February 2026, so nothing in "
        "this database is a legal consent record. communication_opt_in is "
        "Mailchimp subscription status; tracking_consent is a string relayed "
        "from GA4. Use profile_marketing_audience for anything outbound, where "
        "the suppression rules are already applied, and do not present any of "
        "these columns as consent to a compliance question.",
    ),
    (
        "Site-visitor tracking has a start date, and it breaks before-and-after",
        "Anonymous site-visitor tracking came online partway through the "
        "history. Before it, the database held only people who had given an "
        "identifier; after it, every anonymous browser. The resulting step "
        "change is a measurement change, not audience growth. Any comparison "
        "spanning that date is meaningless unless you filter to tier1.",
    ),
    (
        "Count people with COUNT(DISTINCT bn_id) whenever you join",
        "Joins to site_events, content affinity or ad attribution multiply "
        "rows. A bare COUNTIF over one of those joins counts events, not "
        "people. profile_metrics is one row per person, so it is safe on its "
        "own -- the moment you join it to anything event-shaped, switch to "
        "COUNT(DISTINCT bn_id).",
    ),
    (
        "profile_engagement is whole-person lifetime, not per-site",
        "A reader active on three sites contributes their full lifetime totals "
        "to all three rows of any per-site average. Phrase the result as "
        "'engagement level of people who touch this site', never as "
        "'engagement generated on this site'.",
    ),
    (
        "Check fill rates before sizing anything on a sparse field",
        "Several columns exist but are effectively empty: clinical_trials_"
        "interest, treatments_of_interest, symptom_tags, age_band and "
        "caregiver_relationship are all either zero or a few hundred rows. "
        "profile_coverage shows the fill rate for every field. Query it before "
        "building a segment on a column you have not used before.",
    ),
    (
        "Raw email opens are about 84 percent robots",
        "Apple Mail privacy protection and security scanners open messages "
        "automatically. Anything built on email_open_count overstates engagement "
        "roughly six times. is_active_email_90d already uses the bot-filtered "
        "count.",
    ),
]

PROFILE_QUERIES = [
    (
        "1. Known versus unknown, and what we hold on each",
        "The orientation query. Run this before quoting any audience number.",
        """SELECT
    is_known_person,
    COUNT(*)                        AS profiles,
    COUNTIF(has_any_role)           AS with_a_role,
    COUNTIF(condition_key IS NOT NULL) AS with_a_condition,
    COUNTIF(is_mailable)            AS mailable,
    COUNTIF(is_active_email_90d)    AS active_email,
    COUNTIF(has_registration_date)  AS has_real_signup_date
FROM profile_data.profile_metrics
GROUP BY is_known_person
ORDER BY is_known_person DESC""",
    ),
    (
        "2. Audience by condition",
        "Known people, reachable people and clinicians for every condition. This "
        "is the query most requests reduce to.",
        """SELECT
    condition_label,
    COUNTIF(is_known_person)      AS known_people,
    COUNTIF(is_active_email_90d)  AS active_email,
    COUNTIF(is_verified_hcp)      AS hcps,
    COUNTIF(is_patient)           AS patients,
    COUNTIF(is_caregiver)         AS caregivers,
    COUNTIF(is_active_member_90d) AS active_forum_members
FROM profile_data.profile_metrics
WHERE condition_label IS NOT NULL
GROUP BY condition_label
ORDER BY known_people DESC""",
    ),
    (
        "3. Declared or inferred: what a condition count rests on",
        "Run this before quoting any condition audience. Declared sources "
        "(mailchimp_list, site_registration, survey) mean the person told us; "
        "content_affinity means we inferred it from reading behavior at 0.5 "
        "confidence. Only the declared subset supports the word 'patients'. "
        "The source lives on profile_explain, the view built for exactly this "
        "kind of where-did-this-value-come-from question.",
        """SELECT
    x.preferred_condition_source     AS source,
    COUNT(*)                         AS profiles,
    ROUND(AVG(x.preferred_condition_confidence), 2) AS avg_confidence
FROM profile_data.profile_metrics m
JOIN profile_data.profile_explain x USING (bn_id)
WHERE m.condition_key IS NOT NULL
GROUP BY source
ORDER BY profiles DESC""",
    ),
    (
        "4. Role mix within a condition, on an honest denominator",
        "Percentages over people whose role we actually know, not over everyone. "
        "The role_known column is the denominator, and it should be quoted "
        "alongside the percentages.",
        """SELECT
    condition_label,
    COUNTIF(has_any_role)                                             AS role_known,
    ROUND(100 * COUNTIF(is_patient)   / NULLIF(COUNTIF(has_any_role), 0), 1) AS pct_patient,
    ROUND(100 * COUNTIF(is_caregiver) / NULLIF(COUNTIF(has_any_role), 0), 1) AS pct_caregiver,
    ROUND(100 * COUNTIF(is_family_or_friend) / NULLIF(COUNTIF(has_any_role), 0), 1) AS pct_family,
    ROUND(100 * COUNTIF(is_hcp)       / NULLIF(COUNTIF(has_any_role), 0), 1) AS pct_hcp,
    ROUND(100 * COUNTIF(is_other)     / NULLIF(COUNTIF(has_any_role), 0), 1) AS pct_other
FROM profile_data.profile_metrics
WHERE condition_label IS NOT NULL
GROUP BY condition_label
HAVING role_known >= 100
ORDER BY role_known DESC""",
    ),
    (
        "5. Persona mix that sums to 100 percent",
        "Use this when a chart must add up. primary_role is mutually exclusive. "
        "Unspecified is sorted last deliberately -- it is the largest bucket and "
        "leading with it buries the finding.",
        """SELECT
    COALESCE(primary_role, 'unspecified') AS persona,
    COUNT(*)                              AS profiles,
    ROUND(100 * COUNT(*) / SUM(COUNT(*)) OVER (), 1) AS pct
FROM profile_data.profile_metrics
GROUP BY persona
ORDER BY CASE WHEN persona = 'unspecified' THEN 1 ELSE 0 END, profiles DESC""",
    ),
    (
        "6. Who can we actually reach",
        "Mailable means we are allowed to send. Active means they behaved like a "
        "human in the window. The gap between them is the re-engagement "
        "opportunity.",
        """SELECT
    COUNTIF(is_known_person)     AS known_people,
    COUNTIF(is_mailable)         AS mailable,
    COUNTIF(is_active_email_90d) AS active_90d,
    COUNTIF(is_active_email_30d) AS active_30d,
    COUNTIF(is_known_member)     AS forum_members,
    COUNTIF(is_active_member_90d) AS active_members_90d,
    COUNTIF(is_logged_in_90d)    AS logged_in_90d,
    COUNTIF(has_sso_key)         AS has_sso
FROM profile_data.profile_metrics""",
    ),
    (
        "7. Reach by condition and role together",
        "The cross-tab behind most campaign briefs: which clinicians, in which "
        "condition, can we email.",
        """SELECT
    condition_label,
    CASE WHEN is_hcp THEN 'hcp'
         WHEN is_patient THEN 'patient'
         WHEN is_caregiver THEN 'caregiver'
         ELSE 'other or unknown' END AS audience,
    COUNT(*)                     AS people,
    COUNTIF(is_mailable)         AS mailable,
    COUNTIF(is_active_email_90d) AS active_email
FROM profile_data.profile_metrics
WHERE is_known_person AND condition_label IS NOT NULL
GROUP BY condition_label, audience
ORDER BY condition_label, people DESC""",
    ),
    (
        "8. Verified healthcare professionals, and what we hold on them",
        RED + "is_verified_hcp requires real evidence: an NPI number or an "
        "explicit source, and as of 2026-08-19 a credential the registry has not "
        "retired. Presence in the AIM clickstream is not evidence and never sets "
        "the flag. Read the columns left to right -- they narrow from credential "
        "to audience to reachable to live. The drop between the first two is the "
        "one people miss.",
        """SELECT
    condition_label,
    COUNTIF(is_verified_hcp)                         AS verified_hcps,
    COUNTIF(is_engaged_hcp)                          AS engaged_hcps,
    COUNTIF(is_verified_hcp AND is_mailable)         AS reachable,
    COUNTIF(is_verified_hcp AND is_active_email_90d) AS live_90d,
    COUNTIF(has_deactivated_npi)                     AS lapsed_credential
FROM profile_data.profile_metrics
GROUP BY condition_label
HAVING verified_hcps > 0
ORDER BY verified_hcps DESC""",
    ),
    (
        "9. What conditions do our clinicians actually read about",
        "condition_focus looks like the column for this and is permanently empty -- "
        "it is app-reserved and nothing has ever written to it. This is the real "
        "answer, taken from observed reading behaviour in profile_content_affinity, "
        "which covers 57,379 HCPs. affinity_score weights depth of reading rather "
        "than a single visit. Watch how far reachable falls below reading: these are "
        "clinicians we can see but mostly cannot email.",
        """SELECT
    COALESCE(d.label, ca.content_condition)              AS condition_label,
    COUNT(DISTINCT ca.bn_id)                             AS hcps_reading,
    COUNT(DISTINCT IF(m.is_mailable, ca.bn_id, NULL))    AS reachable_hcps,
    ROUND(AVG(ca.affinity_score), 2)                     AS avg_affinity,
    ROUND(AVG(ca.pageview_count), 1)                     AS avg_pageviews
FROM profile_data.profile_content_affinity ca
JOIN profile_data.profile_metrics m USING (bn_id)
LEFT JOIN profile_data.conditions_dict d
       ON LOWER(ca.content_condition) = LOWER(d.condition_key)
       OR LOWER(ca.content_condition) IN UNNEST(ARRAY(SELECT LOWER(a) FROM UNNEST(d.aliases) a))
WHERE m.is_verified_hcp
GROUP BY condition_label
HAVING hcps_reading >= 100
ORDER BY hcps_reading DESC""",
    ),
    (
        "10. Growth: real signups only",
        "Only registration-grade dates are counted. About 88 percent of stored "
        "dates are cookie first-sightings, which are not signups and are heavily "
        "biased toward recent months.",
        """SELECT
    acquisition_month,
    COUNT(*)                     AS new_profiles,
    COUNTIF(is_known_person)     AS new_known_people,
    COUNTIF(is_verified_hcp)     AS new_hcps
FROM profile_data.profile_metrics
WHERE has_registration_date
  AND acquisition_month >= DATE_SUB(DATE_TRUNC(CURRENT_DATE(), MONTH), INTERVAL 24 MONTH)
GROUP BY acquisition_month
ORDER BY acquisition_month DESC""",
    ),
    (
        "11. Audience by registration site",
        "site_domain is the site a person registered on -- filled for only ~7% "
        "of profiles. Useful for per-brand reporting on registrants; for the "
        "full browsing audience use the next query.",
        """SELECT
    site_domain,
    COUNTIF(is_known_person)     AS known_people,
    COUNTIF(is_active_email_90d) AS active_email,
    COUNTIF(is_verified_hcp)     AS hcps,
    COUNT(DISTINCT condition_key) AS conditions_represented
FROM profile_data.profile_metrics
WHERE site_domain IS NOT NULL
GROUP BY site_domain
ORDER BY known_people DESC""",
    ),
    (
        "12. Browsing audience by site: the per-site scorecard",
        "The registration-site query above misses ~94% of people. This one uses "
        "actual browsing (site_events) and shows how contactable, classified and "
        "clinician-heavy each site's audience is. The gap between classified and "
        "contactable is the standing monetization problem.",
        """SELECT e.site_domain,
    COUNT(DISTINCT e.bn_id) AS people,
    ROUND(100 * COUNT(DISTINCT IF(m.is_mailable, e.bn_id, NULL))
          / COUNT(DISTINCT e.bn_id), 1) AS pct_mailable,
    ROUND(100 * COUNT(DISTINCT IF(m.condition_key IS NOT NULL, e.bn_id, NULL))
          / COUNT(DISTINCT e.bn_id), 1) AS pct_classified,
    ROUND(100 * COUNT(DISTINCT IF(m.is_verified_hcp, e.bn_id, NULL))
          / COUNT(DISTINCT e.bn_id), 2) AS pct_clinician
FROM profile_data.site_events e
JOIN profile_data.profile_metrics m USING (bn_id)
GROUP BY e.site_domain
HAVING people >= 1000
ORDER BY people DESC""",
    ),
    (
        "13. Win-back: engaged on email, gone from the site",
        "Mailable people with human email activity in the last 90 days who have "
        "not been seen on the web in that window (or ever), grouped by condition "
        "so the creative can be specific. The IS NULL branch pools never-visited "
        "with lapsed -- treat them differently in the campaign.",
        """SELECT
    m.condition_label,
    COUNT(*) AS winback_candidates
FROM profile_data.profile_metrics m
JOIN profile_data.profile_engagement e USING (bn_id)
WHERE m.is_mailable
  AND m.is_active_email_90d
  AND (e.last_seen_web IS NULL
       OR e.last_seen_web < TIMESTAMP_SUB(CURRENT_TIMESTAMP(), INTERVAL 90 DAY))
GROUP BY m.condition_label
ORDER BY winback_candidates DESC
LIMIT 20""",
    ),
    (
        "14. Which fields are trustworthy enough to build on",
        "The coverage check that should precede any new segment or dashboard. "
        "Sorted worst-first; anything filled in the low single digits cannot "
        "carry a business question. scope is the population each check runs "
        "over (all profiles, the engagement table, or one persona's profiles "
        "for persona-specific fields) -- it is not a persona column.",
        """SELECT
    scope,
    field,
    filled,
    total,
    ROUND(fill_pct, 1) AS fill_pct
FROM profile_data.profile_coverage
ORDER BY scope, fill_pct""",
    ),
    (
        "15. Did the last builds pass their gates?",
        "15 assertions run after every build. Hard failures block publication "
        "entirely, so production keeps serving the last good release; soft "
        "warnings do not. If a number looks wrong, check this before assuming "
        "the query is at fault.",
        """SELECT
    build_id,
    mode,
    status,
    DATE(started_at) AS started,
    assertion_summary
FROM profile_ops.profile_build_runs
ORDER BY started_at DESC
LIMIT 10""",
    ),
    (
        "16. Building a campaign list",
        "The pattern for an actual send. Swap the condition and add filters as "
        "needed. Use profile_current_safe for the identifiers; profile_metrics "
        "carries the flags but not the contact fields.",
        """SELECT
    m.bn_id,
    s.email,
    m.condition_label,
    m.primary_role
FROM profile_data.profile_metrics m
JOIN profile_data.profile_current_safe s USING (bn_id)
WHERE m.is_mailable
  AND m.is_known_person
  AND m.condition_key = 'ms'
  AND m.is_patient""",
    ),
]


# ---------------------------------------------------------------------------
# Identity hub: who someone is, and how confident we are
# ---------------------------------------------------------------------------
HUB_INTRO = [
    "This is the query guide for the identity hub. The hub answers who someone "
    "is; the profile database answers what we know about them. They join on "
    "bn_id.",
    "The hub collapses roughly 34 million identifiers -- emails, cookies, logins, "
    "provider numbers -- into about 7.51 million people. Every query below has "
    "been run against production.",
]

HUB_RULES = [
    (
        "Filter bots out of every people count",
        "Over 200,000 browser identifiers are flagged as automated traffic, and "
        "the flag is advisory -- nothing removes those rows for you. Any people "
        "count without is_bot = FALSE is quietly inflated, and inconsistently so "
        "against reports that did filter.",
    ),
    (
        "Tier is the first filter",
        "tier1 clusters contain something durable: an email address, an SSO key, "
        "a provider number. Those are real people, about 801,000 of them. tier2 "
        "clusters are cookie-only, about 6.7 million, and are closer to browsers "
        "than to humans. Mixing them silently inflates every people number.",
    ),
    (
        "bn_id is stable but not permanent",
        "When new evidence shows two clusters are the same person they merge, and "
        "the losing bn_id is retired and redirected. If you store a bn_id "
        "anywhere, resolve it through bn_id_persistence before using it again. "
        "Splits outnumber merges, so identity is corrected in both directions.",
    ),
    (
        "Always filter bn_id_neighbors by bn_id",
        "It is the cross-product of identifiers inside each cluster and runs to "
        "roughly 75 million rows. Unfiltered it is expensive and rarely what you "
        "wanted.",
    ),
    (
        "bn_id_hub keeps history",
        "Retired edges are marked inactive rather than deleted, so an unfiltered "
        "count includes the past. Add is_active where you want current state.",
    ),
]

HUB_QUERIES = [
    (
        "1. How large is the audience, honestly",
        "Known people versus anonymous browsers, and how many identifiers each "
        "carries.",
        """SELECT
    cluster_tier,
    COUNT(DISTINCT bn_id)                      AS people,
    COUNT(*)                                   AS identifiers,
    ROUND(COUNT(*) / COUNT(DISTINCT bn_id), 2) AS identifiers_per_person
FROM identity_hub_data.bn_id_xref
WHERE is_bot = FALSE
GROUP BY cluster_tier
ORDER BY people DESC""",
    ),
    (
        "2. The identifier-spread health check",
        "How identifiers are spread across people, in five named numbers. The "
        "average alone cannot tell you whether it is propped up by a few monster "
        "clusters; this can. Watch the last column: if it and the 10-plus count "
        "rise together across runs, the graph is starting to glue different "
        "people together (over-merging).",
        """WITH people AS (
    SELECT bn_id, COUNT(*) AS ids
    FROM identity_hub_data.bn_id_xref
    WHERE is_bot = FALSE
    GROUP BY bn_id
),
spread AS (
    SELECT APPROX_QUANTILES(ids, 4) AS q FROM people
)
SELECT
    (SELECT COUNT(*) FROM people)             AS real_people,
    (SELECT SUM(ids) FROM people)             AS total_identifiers,
    q[OFFSET(0)] AS fewest_ids_any_person,
    q[OFFSET(1)] AS ids_bottom_quarter_max,
    q[OFFSET(2)] AS ids_typical_person,
    q[OFFSET(3)] AS ids_top_quarter_min,
    q[OFFSET(4)] AS most_ids_any_person,
    (SELECT COUNTIF(ids >= 10) FROM people)   AS people_with_10_plus_identifiers
FROM spread""",
    ),
    (
        "3. What kinds of identifier do we hold on known people",
        "Reach by identifier type. This is what is actionable.",
        """SELECT
    identifier_type,
    COUNT(DISTINCT bn_id) AS people
FROM identity_hub_data.bn_id_xref
WHERE cluster_tier = 'tier1' AND is_bot = FALSE
GROUP BY identifier_type
ORDER BY people DESC""",
    ),
    (
        "4. Look a person up from any identifier",
        "The everyday operation. Works for an email, a cookie, a provider number "
        "or a login id.",
        """SELECT
    identifier_type,
    identifier_value,
    bn_id,
    cluster_tier,
    cluster_size,
    cluster_health_score
FROM identity_hub_data.bn_id_xref
WHERE identifier_key = CONCAT('email:', LOWER(TRIM('someone@example.com')))""",
    ),
    (
        "5. Everything we know identifies one person",
        "Durable anchors are listed first. Replace the bn_id with a real one.",
        """SELECT
    identifier_type,
    identifier_value,
    last_seen
FROM identity_hub_data.bn_id_xref
WHERE bn_id = 'REPLACE_WITH_BN_ID'
ORDER BY
    CASE identifier_type
        WHEN 'email'      THEN 1
        WHEN 'bionews_uk' THEN 2
        WHEN 'npi_number' THEN 3
        WHEN 'bnfpvid'    THEN 4
        ELSE 9
    END,
    identifier_type""",
    ),
    (
        "6. Why are these two identifiers the same person",
        "The explanation query. match_rule is the reason. Always filter by bn_id.",
        """SELECT
    node_type,
    node_value,
    neighbor_type,
    neighbor_value,
    match_rule,
    source_system,
    ROUND(confidence, 3) AS confidence
FROM identity_hub_data.bn_id_neighbors
WHERE bn_id = 'REPLACE_WITH_BN_ID'
ORDER BY confidence DESC
LIMIT 200""",
    ),
    (
        "7. Cross-device and cross-cookie reach",
        "How many browsers a known person is seen on. This is what identity "
        "resolution buys you: without it each of these would be a separate "
        "'user'.",
        """SELECT
    identifiers_per_person,
    COUNT(*) AS people
FROM (
    SELECT bn_id, COUNTIF(identifier_type IN ('bnfpvid', 'client_id')) AS identifiers_per_person
    FROM identity_hub_data.bn_id_xref
    WHERE cluster_tier = 'tier1' AND is_bot = FALSE
    GROUP BY bn_id
)
WHERE identifiers_per_person > 0
GROUP BY identifiers_per_person
ORDER BY identifiers_per_person""",
    ),
    (
        "8. How much of the email file is visible on the website",
        "A question only the hub can answer: of everyone with an email address, "
        "how many also carry a web identifier -- meaning they can be recognized, "
        "personalized for, and attributed on site. The remainder can be mailed "
        "but is invisible on the web.",
        """WITH email_people AS (
    SELECT DISTINCT bn_id FROM identity_hub_data.bn_id_xref
    WHERE identifier_type = 'email'
),
web_people AS (
    SELECT DISTINCT bn_id FROM identity_hub_data.bn_id_xref
    WHERE identifier_type IN ('bnfpvid', 'client_id')
)
SELECT
    (SELECT COUNT(*) FROM email_people)                                   AS with_email,
    (SELECT COUNT(*) FROM email_people e JOIN web_people w USING (bn_id)) AS also_on_web,
    ROUND(100 * (SELECT COUNT(*) FROM email_people e JOIN web_people w USING (bn_id))
          / (SELECT COUNT(*) FROM email_people), 1)                       AS pct_linked""",
    ),
    (
        "9. Which matching rules build the graph",
        "Read this before trusting or retiring a rule. Only edges at or above "
        "0.80 effective confidence actually merge anyone.",
        """SELECT
    match_rule,
    source_system,
    link_type,
    COUNT(*)                              AS edges,
    COUNT(DISTINCT bn_id)                 AS clusters_touched,
    ROUND(AVG(effective_confidence), 3)   AS avg_confidence,
    COUNTIF(effective_confidence >= 0.80) AS merging_edges
FROM identity_hub_data.bn_id_hub
WHERE is_active
GROUP BY match_rule, source_system, link_type
ORDER BY edges DESC""",
    ),
    (
        "10. Cluster quality and suspicious merges",
        "Oversized clusters are the classic identity failure: a shared device or "
        "a leaked identifier fuses many people into one. Investigate the top of "
        "this list with query 5.",
        """SELECT
    bn_id,
    cluster_tier,
    ANY_VALUE(cluster_size)            AS cluster_size,
    ANY_VALUE(cluster_health_score)    AS health,
    ANY_VALUE(is_shared_workstation)   AS shared_workstation,
    COUNTIF(identifier_type = 'email') AS emails_in_cluster
FROM identity_hub_data.bn_id_xref
GROUP BY bn_id, cluster_tier
HAVING cluster_size > 50
ORDER BY cluster_size DESC
LIMIT 100""",
    ),
    (
        "11. Is the graph still learning: merges and splits by month",
        "Steady monthly volume is a graph that keeps correcting itself; a spike "
        "in merges suggests a new rule is over-merging; a drop to zero means it "
        "stopped learning. Also the audit trail for why a person's id changed.",
        """SELECT
    DATE_TRUNC(event_date, MONTH) AS month,
    event_type,
    COUNT(*)                      AS events,
    SUM(node_count_moved)         AS identifiers_moved
FROM identity_hub_data.bn_id_identity_changes
GROUP BY month, event_type
ORDER BY month DESC, event_type""",
    ),
    (
        "12. Resolve a stored bn_id safely",
        "Run this over any bn_id you saved earlier. If it was retired in a merge "
        "this returns the surviving id.",
        """WITH stored AS (SELECT 'REPLACE_WITH_BN_ID' AS bn_id)
SELECT
    s.bn_id                            AS stored_bn_id,
    COALESCE(p.current_bn_id, s.bn_id) AS use_this_bn_id,
    p.old_bn_id IS NOT NULL            AS was_retired
FROM stored s
LEFT JOIN identity_hub_data.bn_id_persistence p ON p.old_bn_id = s.bn_id""",
    ),
    (
        "13. Activation export: consented segment to ad-platform IDs",
        "The pattern behind every custom-audience or suppression push: the "
        "profile database supplies the consented people, the hub supplies their "
        "Facebook and Google ad identifiers. The consent gate comes entirely "
        "from the marketing view -- never hand-rolled.",
        """SELECT
    m.bn_id,
    x.identifier_type  AS ad_id_type,
    x.identifier_value AS ad_id
FROM profile_data.profile_marketing_audience m
JOIN identity_hub_data.bn_id_xref x
  ON x.bn_id = m.bn_id
 AND x.identifier_type IN ('fbp', 'gcl_au')
WHERE m.is_patient = TRUE
  AND m.preferred_condition.label = 'ms'
LIMIT 100""",
    ),
    (
        "14. Join identity to the profile database",
        "The bridge between the two systems. Identity says who; the profile says "
        "what we know.",
        """SELECT
    x.identifier_type,
    x.identifier_value,
    m.is_known_person,
    m.condition_label,
    m.primary_role,
    m.is_verified_hcp,
    m.is_active_email_90d
FROM identity_hub_data.bn_id_xref x
JOIN profile_data.profile_metrics m USING (bn_id)
WHERE x.identifier_key = CONCAT('email:', LOWER(TRIM('someone@example.com')))""",
    ),
]


def _para(doc, text):
    """Add a paragraph, honouring the RED:: marker."""
    is_red = text.startswith(RED)
    p = doc.add_paragraph()
    run = p.add_run(text[len(RED) :] if is_red else text)
    if is_red:
        run.font.color.rgb = _DARK_RED
    return p


def _mono(par):
    for run in par.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def _build(path: Path, title: str, intro, rules, queries, footer: str) -> None:
    # exports/ holds only generated .docx, which are gitignored, so the
    # directory does not exist in a clean checkout.
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, level=0)

    for para in intro:
        _para(doc, para)

    doc.add_heading("Before you start", level=1)
    for heading, body in rules:
        p = doc.add_paragraph()
        hrun = p.add_run(heading.replace(RED, ""))
        hrun.bold = True
        if heading.startswith(RED):
            hrun.font.color.rgb = _DARK_RED
        _para(doc, body)

    doc.add_heading("Queries", level=1)
    for name, note, sql in queries:
        doc.add_heading(name, level=2)
        _para(doc, note)
        p = doc.add_paragraph(sql)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _mono(p)

    doc.add_heading("Keeping this accurate", level=1)
    doc.add_paragraph(footer)

    doc.save(path)
    print(f"[OK] Wrote {path.relative_to(path.parents[1])}")


def main() -> int:
    _build(
        PROFILE_OUT,
        "Profile Database: Query Guide for BI",
        PROFILE_INTRO,
        PROFILE_RULES,
        PROFILE_QUERIES,
        "Every query in this document is dry-run against production by "
        "tests/unit/test_bi_query_docx.py, so a query that stops working fails "
        "the build rather than reaching an analyst. Regenerate with "
        "python scripts/build_bi_query_docx.py. The runnable versions live in "
        "sql/profile_database_query_cookbook.sql, and the full column reference "
        "is docs/PROFILE_DATABASE_DATA_DICTIONARY.md.",
    )
    _build(
        HUB_OUT,
        "Identity Hub: Query Guide for BI",
        HUB_INTRO,
        HUB_RULES,
        HUB_QUERIES,
        "Every query in this document is dry-run against production by "
        "tests/unit/test_bi_query_docx.py. Regenerate with "
        "python scripts/build_bi_query_docx.py. The runnable versions live in "
        "sql/identity_hub_query_cookbook.sql, and the full column reference is "
        "docs/IDENTITY_HUB_DATA_DICTIONARY.md.",
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
