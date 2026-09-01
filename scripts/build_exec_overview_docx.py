#!/usr/bin/env python3
"""
Build exports/profile_database_executive_overview.docx.

Executive overview of the Profile Database and the Identity Hub, in three
sections: what each one does and delivers, then how they work together.

Written for executive management. No technology detail, no schema, no
operational content -- the test for every sentence is whether someone outside
the data team can read it once and understand it.

Every figure is production-verified 2026-08-24.

  python scripts/build_exec_overview_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
EXPORTS = REPO / "exports"
OUT = EXPORTS / "profile_database_executive_overview.docx"

NAVY = RGBColor(0x0D, 0x47, 0xA1)
GREY = RGBColor(0x5F, 0x63, 0x68)


def para(doc, text, size=11, bold=False, italic=False, color=None, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def lead(doc, label, text):
    """A named benefit: bold lead-in, then the explanation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(label + "  ")
    r.font.size = Pt(11)
    r.font.bold = True
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = NAVY
    return h


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p


def table(doc, rows, headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light List Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        if i > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(10)
            if i > 0:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t


def build():
    doc = Document()

    doc.add_heading("Knowing Our Audience", level=0)
    para(
        doc,
        "The Profile Database and the Identity Hub",
        size=13,
        italic=True,
        color=GREY,
    )
    para(doc, "Executive overview  |  August 2026", size=10, color=GREY, after=20)

    para(
        doc,
        "BioNews now knows who its audience is. Not approximately, and not as a monthly "
        "estimate somebody assembles by hand, but person by person, refreshed every "
        "morning, and precise enough to put in front of a client.",
    )
    para(
        doc,
        "Two systems make that possible. The Identity Hub works out who somebody is. The "
        "Profile Database records what we know about them. This document covers what each "
        "one delivers, and then what they make possible together.",
    )

    # ── SECTION 1 ────────────────────────────────────────────────────────────
    h1(doc, "Part One: The Profile Database")

    para(
        doc,
        "The Profile Database is a single record for every person in our audience. "
        "Everything we know about somebody, in one place, kept current.",
    )
    para(
        doc,
        "Before it existed, the answer to a simple question like how many patients do we "
        "have for this condition lived in five or six different systems that disagreed "
        "with one another. Getting an answer meant somebody exporting spreadsheets and "
        "reconciling them by hand, and the answer was out of date by the time it was "
        "finished. Now it is one question, asked once, and the answer is right.",
    )

    h2(doc, "What we hold today")
    table(
        doc,
        [
            ("Total profiles", "7,495,014"),
            ("People we can name and contact", "798,780"),
            ("People whose role we know", "499,187"),
            ("Reachable by email, with consent", "407,802"),
            ("Engaged with email in the last 90 days", "188,196"),
            ("Conditions covered", "87"),
            ("Sites represented", "94"),
        ],
        ("Measure", "Today"),
    )

    h2(doc, "What it does")

    lead(
        doc,
        "One record per person.",
        "Everything we know about somebody sits together: their condition, their role, "
        "how they found us, what they read, what they have agreed to receive, and when we "
        "last heard from them.",
    )
    lead(
        doc,
        "It knows who it is talking to.",
        "For 499,187 people we know whether they are a patient, a caregiver, a family "
        "member or a clinician. That one fact changes how everything should be written, "
        "targeted and sold, and until now we were largely guessing at it.",
    )
    lead(
        doc,
        "It separates what we were told from what we worked out.",
        "3.2 million profiles carry a condition. For every one, we know whether the person "
        "told us directly or whether we inferred it from what they read. That is the "
        "difference between saying somebody is a patient and saying somebody is "
        "interested, and being able to make it protects us commercially and ethically.",
    )
    lead(
        doc,
        "It verifies our clinicians.",
        "380,508 healthcare professionals with a real credential we have checked against "
        "the federal provider registry, including their specialty and where they practise. "
        "Not somebody who ticked a box on a form.",
    )
    lead(
        doc,
        "It knows who we are allowed to contact.",
        "Consent and subscription status are part of the record, not a separate list "
        "somebody remembers to check. When we say we can reach 407,802 people, that "
        "number already accounts for permission.",
    )
    lead(
        doc,
        "It tracks engagement, not just presence.",
        "We know who opens, who clicks, who visits, who posts in the forum, and who has "
        "gone quiet. An audience of people who actually pay attention is worth far more "
        "than a list.",
    )
    lead(
        doc,
        "It rebuilds itself every morning.",
        "Nobody is working from a file somebody exported in March. The whole thing is "
        "current daily and checks itself before anyone sees it.",
    )
    lead(
        doc,
        "It comes with the answers already built.",
        "Common audiences are pre-built, so the business is not re-inventing a definition "
        "every time somebody asks a question and getting a slightly different number.",
    )
    lead(
        doc,
        "It shows us what we do not know.",
        "The system reports how complete it is, field by field. That turns vague ambitions "
        "into a specific list of things worth going out and collecting.",
    )

    h2(doc, "What this gives BioNews")

    lead(
        doc,
        "Numbers we can defend.",
        "Every audience figure can be traced back to why we believe it. When a partner "
        "asks a second question, we have an answer.",
    )
    lead(
        doc,
        "We stop contacting the same person twice.",
        "Because we know three records are one human being, we speak to them once, in the "
        "right voice, about the condition they actually care about. That protects the "
        "relationship we have with our readers.",
    )
    lead(
        doc,
        "We can size an opportunity before committing to it.",
        "If we are considering a new site or a new condition, we can say how many people "
        "we already have who would care on day one. That is a business case built on "
        "people we own rather than a market estimate.",
    )
    lead(
        doc,
        "We can price and sell honestly.",
        "For any condition we can produce the total audience, the reachable subset and the "
        "actively engaged subset, and stand behind all three.",
    )
    lead(
        doc,
        "We can find audiences we did not know we had.",
        "Because everything is in one place, patterns surface that nobody thought to look "
        "for, such as which conditions our clinicians read about as opposed to the ones we "
        "assumed they cared about.",
    )
    lead(
        doc,
        "Answers in seconds instead of days.",
        "Questions that used to require a person and a week are now self-service.",
    )

    h2(doc, "The questions it answers")
    for q in [
        "How many diagnosed patients do we have for a condition, and how many told us "
        "directly rather than us inferring it.",
        "How many people can we email tomorrow, with consent, for a specific condition, "
        "and how many of those actually open anything.",
        "Which conditions are our clinicians genuinely reading about.",
        "Which of our 94 sites bring in real registrations rather than just traffic.",
        "How is a condition audience growing month over month, and is that growth real "
        "people or repeat visits.",
        "Who joined the forum but never opens an email, and who reads every email but has "
        "never joined anything.",
        "Which audiences are large enough to build a campaign around, and which only look "
        "large until you check who we can actually reach.",
        "If we launched a site for a new condition tomorrow, how many people do we already "
        "have who would be interested.",
    ]:
        bullet(doc, q)

    h2(doc, "What that looks like in practice")
    para(
        doc,
        "Our eight largest condition audiences, with the number of people we can name and "
        "the number we can actually reach. We can produce this for all 87 conditions on "
        "demand, and the reachable column is the one we can sell against.",
    )
    table(
        doc,
        [
            ("Multiple Sclerosis", "59,361", "53,725"),
            ("Parkinsons Disease", "58,520", "51,977"),
            ("ALS", "24,640", "21,309"),
            ("Pulmonary Fibrosis", "20,477", "15,732"),
            ("Ehlers-Danlos Syndrome", "17,048", "16,231"),
            ("Sjogrens Syndrome", "16,516", "15,811"),
            ("Cystic Fibrosis", "15,725", "14,723"),
            ("Muscular Dystrophy", "15,355", "14,102"),
        ],
        ("Condition", "Known people", "Reachable"),
    )

    h2(doc, "The clinician audience")
    para(
        doc,
        "This deserves separate attention because it is a genuine commercial asset and very "
        "few publishers in our space have anything comparable.",
    )
    table(
        doc,
        [
            ("Verified clinicians", "380,508"),
            ("Who have engaged with our content", "102,584"),
            ("Who we can reach by email", "23,027"),
            ("Active on email in the last 90 days", "13,937"),
        ],
        ("Healthcare professionals", "Today"),
    )
    para(
        doc,
        "We track four numbers rather than one because they answer four different "
        "questions, and using the wrong one is how organisations end up promising "
        "something they cannot deliver. The system makes the right number the easy one to "
        "reach for, which is exactly the discipline you want when the answer is going into "
        "a client conversation.",
    )

    # ── SECTION 2 ────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Part Two: The Identity Hub")

    para(
        doc,
        "The Identity Hub answers one question: who is this person. It is the part nobody "
        "sees, and nothing else works without it.",
    )
    para(
        doc,
        "One person arrives at BioNews as many different records. An email address, a "
        "browser on a laptop, another on a phone, a newsletter subscription, a forum "
        "login, a form they filled in, a click on an advert. Left alone, that one person "
        "looks like seven people. Every audience number is inflated, the same person gets "
        "emailed repeatedly, and nobody can say by how much any of it is wrong.",
    )
    para(
        doc,
        "The Identity Hub takes 34.6 million of these separate identifiers and works out "
        "which of them belong to the same human being. It weighs 75.8 million individual "
        "pieces of evidence to reach those conclusions.",
    )

    h2(doc, "What it does")

    lead(
        doc,
        "It turns fragments into people.",
        "34.6 million scattered identifiers become a single, durable identity for each "
        "person, which everything else in the business can rely on.",
    )
    lead(
        doc,
        "It keeps its reasoning.",
        "Every piece of evidence it considered is retained, including the evidence it "
        "rejected. Any conclusion it reached can be explained rather than simply trusted.",
    )
    lead(
        doc,
        "It separates certain from probable.",
        "A shared email address is proof. Two visits from the same device is a hint. The "
        "Hub treats those differently and only acts on evidence that clears a defined bar.",
    )
    lead(
        doc,
        "It corrects itself as it learns.",
        "When new evidence shows that two people are actually one, it merges them, and "
        "everything downstream follows automatically without breaking.",
    )
    lead(
        doc,
        "It filters out what is not human.",
        "Automated traffic is identified and excluded, so our audience numbers reflect "
        "people rather than machines.",
    )
    lead(
        doc,
        "It remembers people for thirteen months.",
        "Somebody who reads about a condition, disappears for eight months and comes back "
        "is still recognised as the same person. This matters more in our audience than "
        "almost any other, because people engage around diagnosis and treatment rather "
        "than continuously. 4.5 million of the people we currently recognise were last "
        "seen more than ninety days ago, and 529,890 of them more than a year ago. A "
        "shorter memory would simply have forgotten them.",
    )
    lead(
        doc,
        "It is honest about what it does not know.",
        "It clearly separates the people we can name from the visitors we merely "
        "recognise, so nobody accidentally quotes the larger number.",
    )

    h2(doc, "The thirteen month memory")

    para(
        doc,
        "This is worth setting out on its own, because it is the single change that has "
        "most increased what we know.",
    )
    para(
        doc,
        "Identity systems have to decide how long to remember somebody who has gone quiet. "
        "Remember for too short a period and you forget people who were always coming back. "
        "Our audience is not a retail audience that visits every week. Somebody researches "
        "a diagnosis intensively for a month, goes quiet while they get on with treatment, "
        "and returns when something changes. A system that forgets them in ninety days "
        "treats every one of those returns as a brand new stranger.",
    )
    para(
        doc,
        "We now remember for thirteen months. The effect is substantial:",
    )
    table(
        doc,
        [
            ("People we currently recognise", "7,522,950"),
            ("Last seen more than 90 days ago", "4,508,886"),
            ("Last seen more than 180 days ago", "1,718,142"),
            ("Last seen more than a year ago", "529,890"),
        ],
        ("Reach of the thirteen month memory", "People"),
    )
    para(
        doc,
        "4.5 million of the people we can recognise today were last seen more than ninety "
        "days ago. Under a shorter memory they would be strangers to us. Instead, when any "
        "one of them returns, we already know their condition, their role, what they read "
        "and what they have agreed to receive.",
    )

    h2(doc, "Why ours is better than most")

    para(
        doc,
        "Almost every company is building something like this now, so it is fair to ask "
        "what makes ours worth having.",
    )

    lead(
        doc,
        "It is ours.",
        "Most companies license identity from an outside vendor, which means renting a "
        "third party's view of their own audience. Ours is built entirely from our own "
        "sites, newsletters, forums and registrations. Nobody can take it away, raise the "
        "price, change the terms, or sell the same picture to a competitor.",
    )
    lead(
        doc,
        "It can be explained.",
        "Most identity products are a black box that hands back an answer. Ours can show "
        "why it believes two records are the same person. In healthcare, where we may have "
        "to justify that judgement, this is not a nice to have.",
    )
    lead(
        doc,
        "It knows our audience, not a generic one.",
        "A bought identity graph is built for retail shoppers and adapted to whatever you "
        "do. Ours was designed for exactly one job: knowing patients, caregivers and "
        "clinicians in rare disease. It understands conditions, roles and professional "
        "credentials because it was built to.",
    )
    lead(
        doc,
        "It carries verified professional credentials.",
        "380,508 clinicians checked against the federal registry, with specialty and "
        "practice location. No general purpose vendor graph has this, and it would be "
        "expensive to buy.",
    )
    lead(
        doc,
        "It lives where the rest of our data lives.",
        "Many identity systems require a separate specialist platform, with its own "
        "licence, its own skills and a permanent translation layer between it and "
        "everything else. Ours sits alongside the rest of the business, so anyone who can "
        "already ask a question of our data can use it. There is no second system to buy "
        "or staff.",
    )
    lead(
        doc,
        "It was built by someone who has done it before.",
        "This is the difference between a two year science project and a working system. "
        "BioNews got the benefit of knowing in advance where these builds usually go "
        "wrong.",
    )

    h2(doc, "What this gives BioNews")

    lead(
        doc,
        "We count people, not devices.",
        "Every audience figure the business quotes is now a count of human beings. That "
        "single change makes every downstream number trustworthy.",
    )
    lead(
        doc,
        "We own our audience outright.",
        "No vendor dependency, no renewal risk, no third party with the same picture of "
        "our readers.",
    )
    lead(
        doc,
        "It gets better with everything we add.",
        "Every new source, site or platform is matched against people we already know. "
        "The value compounds rather than simply accumulating.",
    )
    lead(
        doc,
        "It protects the reader relationship.",
        "Recognising somebody properly means not asking them the same question twice and "
        "not sending them the same message three times.",
    )
    lead(
        doc,
        "It is an asset on the balance sheet, not a subscription.",
        "What we have built has lasting value that grows. A licence is a cost that recurs.",
    )

    # ── SECTION 3 ────────────────────────────────────────────────────────────
    doc.add_page_break()
    h1(doc, "Part Three: How They Work Together")

    para(
        doc,
        "The simplest way to picture it is a switchboard and a file. The Identity Hub "
        "answers the question of who is calling. The Profile Database is the file it hands "
        "you once it knows. Neither is much use alone. Together they are the reason we can "
        "say anything confident about our audience at all.",
    )

    para(
        doc,
        "This is where most companies come unstuck. Plenty of organisations build an "
        "identity graph, then discover that on its own it answers nothing. It tells you "
        "that seven records are one person, which is interesting but not actionable. Ours "
        "was designed from the start to feed something, and the Profile Database was "
        "designed from the start to be fed. They were built as one system with two jobs.",
    )

    h2(doc, "What that combination unlocks")

    para(
        doc,
        "The thirteen month memory is a good illustration of how the two work as one. The "
        "Identity Hub established it first, and the Profile Database has been steadily "
        "absorbing it ever since, adding people it can now recognise. The hub currently "
        "recognises 7,522,950 people and the Profile Database holds 7,495,014 of them, "
        "with the remainder flowing through on the next daily rebuild. The gap is small "
        "and it closes itself, which is exactly how it should work: the Hub establishes "
        "who exists, and the Profile Database follows.",
    )

    lead(
        doc,
        "Every source makes every profile richer.",
        "When we connect something new, the Hub works out how those identifiers relate to "
        "people we already know, and the Profile Database deepens automatically. We are "
        "not starting over each time. We are adding to something that compounds.",
    )
    lead(
        doc,
        "A single view across 94 sites.",
        "Somebody who reads two of our publications is one person to us, not two audiences. "
        "We can see the whole relationship rather than a fragment of it.",
    )
    lead(
        doc,
        "Behaviour connected to identity.",
        "Knowing what somebody read is only useful if you know who they are. Knowing who "
        "they are is only useful if you know what they care about. The combination is what "
        "makes either worth having.",
    )
    lead(
        doc,
        "Growth we can measure honestly.",
        "We can tell genuine new audience from the same people returning, which means we "
        "know whether we are actually growing.",
    )
    lead(
        doc,
        "Confidence when it counts.",
        "When we make a claim to a partner, a client or a board, it is backed by a named, "
        "consented, verifiable audience.",
    )

    h2(doc, "Where this goes next")

    lead(
        doc,
        "Our own registration becomes the front door.",
        "SurveyEngine, our own platform, is becoming the primary way people sign up. This "
        "is the single biggest improvement ahead, because when somebody registers with us "
        "directly we learn far more about them, and we learn it because they told us "
        "rather than because we worked it out. As it becomes the main route in, the "
        "quality of what we know rises across the whole audience.",
    )
    lead(
        doc,
        "More sources, chosen deliberately.",
        "We can now see precisely which gaps are worth closing, so each new connection is "
        "a decision rather than a guess. And because everything is matched against people "
        "we already have, each one deepens what we know rather than just adding names.",
    )
    lead(
        doc,
        "Deeper clinician intelligence.",
        "The professional audience carries the most commercial interest, and we are "
        "continuing to widen what we can say about it.",
    )
    lead(
        doc,
        "Self service across the business.",
        "The direction of travel is that anyone who needs an audience number can get it "
        "themselves, correctly, without waiting on anyone.",
    )

    h2(doc, "In short")
    para(
        doc,
        "BioNews has gone from describing its audience in approximations to knowing it "
        "person by person, refreshed every morning. We can size any condition, identify any "
        "role, verify any clinician, and reach anyone who has given us permission. We can "
        "tell the difference between what we know and what we have inferred, which lets us "
        "be ambitious and accurate at the same time.",
    )
    para(
        doc,
        "We own it outright, it was built for our business rather than adapted from "
        "somebody else's, and it gets more valuable with every source we connect. That is "
        "a strong position for a publisher in this space, and it is a durable one.",
    )

    EXPORTS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    words = sum(len(p.text.split()) for p in doc.paragraphs) + sum(
        len(c.text.split()) for t in doc.tables for r in t.rows for c in r.cells
    )
    print(
        f"[OK] Wrote {OUT.relative_to(REPO)}  ({words:,} words, {len(doc.tables)} tables)"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(build())
