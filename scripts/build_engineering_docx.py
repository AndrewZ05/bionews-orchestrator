#!/usr/bin/env python3
"""
Build the two engineering-audience .docx documents:

  exports/identity_hub_engineering.docx
  exports/profile_db_engineering.docx

Audience is engineers, not BI or management. These explain how the two systems
are actually built -- the algorithms, the execution model, the failure modes and
the technology choices -- and defer every column-level question to the data
dictionaries rather than repeating them.

Deliberately NOT a data dictionary and NOT a runbook. The dictionaries are
generated from the DDL and stay correct on their own; the runbooks cover
operating procedure. This pair covers the part that lives only in the code:
why it works the way it does.

  python scripts/build_engineering_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
EXPORTS = REPO / "exports"
HUB_OUT = EXPORTS / "identity_hub_engineering.docx"
PROFILE_OUT = EXPORTS / "profile_db_engineering.docx"
AS_OF = "24 August 2026"

MONO = "Consolas"
DARK_RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x5F, 0x63, 0x68)


# ── rendering helpers ─────────────────────────────────────────────────────────


def para(doc, text, size=10.5, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def code(doc, text):
    """Monospace block. Used for signatures, SQL fragments and file paths."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = MONO
    r.font.size = Pt(9)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


def kv_table(doc, rows, headers=("", "")):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            if i == 0 and any(c in str(val) for c in "._(") and " " not in str(val):
                r.font.name = MONO
    return t


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def preamble(doc, title, subtitle, scope_note):
    doc.add_heading(title, level=0)
    para(doc, subtitle, size=11.5, italic=True, color=GREY)
    para(doc, scope_note, size=10)
    para(
        doc, f"Verified against production {AS_OF}.", size=9.5, italic=True, color=GREY
    )


# ── identity hub ──────────────────────────────────────────────────────────────


def build_hub():
    doc = Document()
    preamble(
        doc,
        "Identity Hub -- Engineering Reference",
        "How the graph is built: algorithms, execution model, and the failure modes worth knowing about.",
        "Audience: engineers. Column-level questions are answered by "
        "identity_hub_data_dictionary.docx, which is generated from the DDL and the "
        "pipeline config; this document does not repeat it. Operating procedure for "
        "a full rebuild is in docs/IDENTITY_HUB_FULL_REBUILD_RUNBOOK.md.",
    )

    h1(doc, "1. What the system is")
    para(
        doc,
        "An entity-resolution pipeline. It consumes identifiers from 23 enabled "
        "connectors, decides which of them belong to the same person, and assigns "
        "each resulting cluster a durable bn_id. Everything downstream -- the "
        "profile database, audience exports, ad-platform activation -- joins on that "
        "bn_id.",
    )
    para(
        doc,
        "It is BigQuery-to-BigQuery. Unlike the other extractors in this repo there "
        "is no Parquet, no GCS staging and no external API. The implementation lives "
        "in shared/identity_hub.py (~10,600 lines) and runs in-process under "
        "plugins/identity_hub_extractor.py, which exists mainly to give the "
        "orchestrator a standard run_pipeline() entry point and environment routing.",
    )
    code(
        doc,
        "python orchestrate.py --source identity_hub --env prod --refresh full\n"
        "python orchestrate.py --source identity_hub --env prod --lookback 7\n"
        "python orchestrate.py --source identity_hub --env prod --refresh full --dry-run",
    )
    para(
        doc,
        "Note the layering: schema_prefix and schema_suffix are accepted by the "
        "orchestrator but affix staging table names in the Parquet path, which this "
        "pipeline never touches. Passing them once looked like sandboxing and "
        "actually rebuilt production. They now raise (commit 1f7a170).",
        color=DARK_RED,
    )

    h1(doc, "2. Technology")
    kv_table(
        doc,
        [
            (
                "BigQuery",
                "Storage and the entire edge-processing layer. Scoring, decay, quality filters and gates are SQL, not Python.",
            ),
            (
                "Python 3.13",
                "Orchestration, Union-Find, and bn_id derivation. The only stage that materialises data in process memory.",
            ),
            (
                "google-cloud-bigquery",
                "Client library. Query jobs are labelled per step for cost attribution.",
            ),
            (
                "SHA-256 + base64url",
                "bn_id derivation. Deterministic, so identical cluster membership always yields an identical bn_id.",
            ),
            (
                "YAML",
                "configs/identity_hub.yaml is the whole behavioural contract: connectors, thresholds, decay bands, priorities.",
            ),
        ],
        headers=("Component", "Role"),
    )

    h1(doc, "3. Execution model")
    para(
        doc,
        "A run has five phases. Phases 1-3 are SQL over a staging table; only phase 4 "
        "pulls data into Python.",
    )
    kv_table(
        doc,
        [
            (
                "1. Connector fan-in",
                "Each enabled connector emits candidate edges into a per-run staging table via _insert_connector_edges(). Batched, never held in memory.",
            ),
            (
                "2. Aggregation",
                "_run_bq_aggregation() collapses duplicate edges and combines evidence into a single confidence per identifier pair.",
            ),
            (
                "3. Filter and gate",
                "_run_bq_quality_filters() and _run_bq_gates() apply fanout caps, shared-workstation detection, bot exclusion and per-rule confidence caps.",
            ),
            (
                "4. Union-Find",
                "_export_union_find_tuples() streams surviving edges into PriorityUnionFind; components become clusters.",
            ),
            (
                "5. Write and verify",
                "_write_hub_from_staging(), then _verify_output_contracts() and the shrink safeguard before anything replaces production.",
            ),
        ],
        headers=("Phase", "What happens"),
    )

    h2(doc, "3.1 Union-Find, and why it is priority-weighted")
    para(
        doc,
        "PriorityUnionFind (shared/identity_hub.py:158) is a standard disjoint-set "
        "structure with path compression and union by rank, plus one addition: every "
        "node carries a source priority, and the lower number wins the root.",
    )
    code(
        doc,
        "if pri_x > pri_y:\n"
        "    root_x, root_y = root_y, root_x\n"
        "elif pri_x == pri_y and root_x > root_y:\n"
        "    root_x, root_y = root_y, root_x",
    )
    para(
        doc,
        "That tie-break on the node value itself is not cosmetic. It makes root "
        "selection independent of the order edges arrive in, which is what makes a "
        "rebuild reproducible: the same edge set always produces the same roots, "
        "therefore the same bn_ids.",
    )
    para(
        doc,
        "Priority is why an email or an SSO key ends up as the canonical root rather "
        "than a cookie. The root is then hashed:",
    )
    code(doc, 'bn_id = "BN_" + base64url(SHA256(canonical_root))[:16]')
    para(
        doc,
        "Consequence worth internalising: bn_id is a pure function of cluster "
        "membership. Nothing is allocated or stored. Two clusters merging produces a "
        "new root and therefore a new bn_id, and the loser is recorded in "
        "bn_id_persistence. Any bn_id held outside the system must be resolved "
        "through that table before use.",
    )

    h2(doc, "3.2 Confidence, and the single number that decides")
    para(
        doc,
        "Three confidence columns exist and only one of them decides anything. "
        "base_confidence is what the connector asserted. confidence is the aggregate "
        "after combining repeated evidence. effective_confidence is that aggregate "
        "after time decay and the per-rule cap -- and it alone is compared against "
        "stitch_threshold (0.80).",
    )
    para(
        doc,
        "Edges below the threshold are still written. bn_id_hub is an evidence log, "
        "not a decision log: it holds 75.4M edges including everything rejected. "
        "Counting rows in it does not count linkages.",
    )

    h1(doc, "4. Configuration is the contract")
    para(
        doc,
        "23 of 27 declared connectors are enabled: 18 deterministic, 5 "
        "probabilistic. Eight are marked static: true, meaning they are reference "
        "data whose edges already exist in the prior graph and are skipped on "
        "incremental runs.",
    )
    kv_table(
        doc,
        [
            (
                "stitch_threshold",
                "0.80",
                "Minimum effective_confidence to merge two identifiers.",
            ),
            (
                "max_visitors_per_identifier",
                "125",
                "Fanout cap. Above this an identifier is treated as shared or invalid and excluded.",
            ),
            (
                "shrink_abort_threshold",
                "0.90",
                "Refuse to overwrite an output table that came back below 90% of its previous row count.",
            ),
            (
                "decay_schedule",
                "0.8 band = 400d",
                "13-month retention. 1.0 x 0.8 = 0.80 exactly meets the threshold; the next band falls below and is the intended cutoff.",
            ),
        ],
        headers=("Setting", "Value", "Meaning"),
    )
    para(
        doc,
        "Every run stamps config_version (a hash of the resolved config) and git_sha "
        "into bn_id_metrics. Two runs that disagree are attributable to one or the "
        "other without guesswork.",
    )

    h1(doc, "5. Incremental versus full, and why it matters")
    para(
        doc,
        "The default run is incremental: it appends new edges and attaches them to "
        "existing clusters using bn_id_node_index, the previous run's answer key. It "
        "does not re-evaluate old edges. A full rebuild re-derives every cluster from "
        "the entire edge history.",
    )
    para(
        doc,
        "That difference is the single most important operational fact about this "
        "system. Configuration changes affecting decay, influence windows or edge "
        "expiry have NO effect on an incremental run. They take effect, all at once, "
        "on the next full rebuild -- which may be months later, under a different "
        "engineer, with the change long forgotten.",
        bold=True,
    )
    para(
        doc,
        "This is not hypothetical. On 2026-08-06 a full rebuild aged out 2,085,967 "
        "anonymous identities under a 90-day effective lifetime that had emerged from "
        "three interacting settings and had never been approved as a rule. It was "
        "invisible for months precisely because incrementals never re-evaluate old "
        "edges. Retention is now 400 days (commit 3735ee9), and that fix has itself "
        "not yet been exercised by a full rebuild.",
        color=DARK_RED,
    )

    h1(doc, "6. Safety mechanisms")
    kv_table(
        doc,
        [
            (
                "Shadow write",
                "A full rebuild writes to a shadow table and only swaps after validation, so a failure mid-run leaves production intact.",
            ),
            (
                "Shrink safeguard",
                "_check_shrink_safeguard() aborts the write when a table returns below shrink_abort_threshold of its previous size. Bypassed only by --force-overwrite.",
            ),
            (
                "Output contracts",
                "_verify_output_contracts() checks row counts and referential shape before promotion.",
            ),
            (
                "Manifest promotion",
                "bn_id_manifest records which run is live. Consumers read the latest promoted_at row; a half-written run is never visible.",
            ),
            (
                "Connector failure gate",
                "_raise_on_connector_failures() refuses to promote a run where a connector errored, rather than silently publishing a partial graph.",
            ),
        ],
        headers=("Mechanism", "What it protects against"),
    )
    para(
        doc,
        "The shrink safeguard was 0.5 until 2026-08-21. The 2026-08-06 truncation "
        "arrived at 57.9% of the previous row count -- above the threshold, so it "
        "wrote cleanly. It is now 0.90 (commit e59f557).",
    )

    h1(doc, "7. Failure modes seen in production")
    kv_table(
        doc,
        [
            (
                "Silent config-driven eviction",
                "A rebuild applies months of accumulated config change at once. Mitigated by the 0.90 shrink threshold and the downstream identity_source_row_delta gate.",
            ),
            (
                "Under-instrumented rebuilds",
                "A healthy run writes 96 rows to bn_id_metrics. The run that caused the 8/07 damage wrote none. Absence of a bad signal is not evidence of a good run.",
            ),
            (
                "Unaliased UNION branch",
                "A dry run validates that SQL parses, not that output is sane. This class of bug has reached production through a passing preflight.",
            ),
            (
                "Stale consumer bn_id",
                "Held bn_ids go stale when clusters merge. Always resolve through bn_id_persistence.",
            ),
        ],
        headers=("Failure mode", "Detail"),
    )

    h1(doc, "8. Where to look next")
    kv_table(
        doc,
        [
            (
                "shared/identity_hub.py",
                "The implementation. PriorityUnionFind at :158, IdentityHubBuilder at :281.",
            ),
            (
                "configs/identity_hub.yaml",
                "Connectors, thresholds, decay bands, type priorities.",
            ),
            (
                "identity_hub_data_dictionary.docx",
                "Every table and column. Generated -- if it disagrees with the database, regenerate it.",
            ),
            (
                "identity_hub_table_reference.docx",
                "Confidence and cluster-health scoring written out in full.",
            ),
            (
                "docs/IDENTITY_HUB_FULL_REBUILD_RUNBOOK.md",
                "Pre-rebuild snapshot, stop rule, restore path.",
            ),
        ],
        headers=("Path", "What it holds"),
    )

    EXPORTS.mkdir(parents=True, exist_ok=True)
    doc.save(HUB_OUT)
    print(f"[OK] Wrote {HUB_OUT.relative_to(REPO)}")


# ── profile database ──────────────────────────────────────────────────────────


def build_profile():
    doc = Document()
    preamble(
        doc,
        "Profile Database -- Engineering Reference",
        "How the build works: modes, step manifest, publication gates, and the invariants that keep a bad build off production.",
        "Audience: engineers. Column-level questions are answered by "
        "profile_db_data_dictionary.docx, which is generated from the DDL; this "
        "document does not repeat it. Day-to-day operating procedure is in "
        "docs/PROFILE_DATABASE_OPERATOR_GUIDE.md.",
    )

    h1(doc, "1. What the system is")
    para(
        doc,
        "A SQL-only ELT pipeline. It takes the identity graph as its input contract "
        "and merges eight source systems onto it, producing one row per person in "
        "profile_core plus a set of satellite tables and consumer views.",
    )
    para(
        doc,
        "There is no extraction stage. plugins/profile_database_extractor.py (~5,200 "
        "lines) is an orchestrator over ordered .sql files: it resolves the step list "
        "for a mode, executes each file, records timing and row counts, runs "
        "assertions, and decides whether the result may be published.",
    )

    h1(doc, "2. Technology")
    kv_table(
        doc,
        [
            (
                "BigQuery",
                "All transformation. Every step is a .sql file executed against production or a candidate dataset.",
            ),
            (
                "Python 3.13",
                "Step sequencing, telemetry, assertions, publication logic. No row-level data passes through Python.",
            ),
            (
                "shared/profile_database_manifest.py",
                "The single source of truth for steps and modes: 36 declared steps, mapped into 7 modes.",
            ),
            (
                "profile_ops dataset",
                "Build telemetry: profile_build_runs, profile_build_steps, profile_publish_manifest, profile_dataset_leases.",
            ),
            (
                "BigQuery table snapshots",
                "Monthly forensic copies, 13-month retention, via scripts/snapshot_monthly_identity_profile.py.",
            ),
        ],
        headers=("Component", "Role"),
    )

    h1(doc, "3. Build modes")
    para(
        doc,
        "A mode is an ordered subset of the 36 declared steps. The step list is data, "
        "not control flow, which is what makes resume modes possible.",
    )
    kv_table(
        doc,
        [
            (
                "rebuild",
                "33",
                "Full reconstruction. Drops and repopulates profile_core from the identity graph.",
            ),
            (
                "refresh",
                "22",
                "The daily run. Scoped to changed profiles; does not repopulate.",
            ),
            (
                "reenrich",
                "18",
                "Enrichment, personas and views only. No DDL, no populate.",
            ),
            ("resume_rebuild", "3", "Continue a rebuild that failed after populate."),
            ("resume_publish", "0", "Publish an already-built candidate dataset."),
            ("views", "1", "Republish consumer views only."),
            ("backfill_site_events", "1", "Targeted historical load."),
        ],
        headers=("Mode", "Steps", "Purpose"),
    )
    para(
        doc,
        "14 of the 33 rebuild steps never run in refresh -- all four fill_gaps_* "
        "steps, every populate_*, the DDL, snapshot and restore. Any column written "
        "only by those steps is frozen at the last rebuild. This is a recurring source "
        "of silent staleness: npi_deactivation_date drifted this way and was moved "
        "into the maintenance step (commit b2134bc) precisely because a one-off "
        "backfill fixes a symptom on a schedule of never.",
        color=DARK_RED,
    )

    h1(doc, "4. Scope, and the predicate audit")
    para(
        doc,
        "A refresh operates on a scope set -- the bn_ids with new activity, typically "
        "~400K of 7.5M. Every enrichment statement must either carry a scope "
        "predicate or be explicitly declared a global reconciliation.",
    )
    code(doc, "-- SCOPE: global_reconcile\nMERGE INTO profile_data.profile_core pc ...")
    para(
        doc,
        "scripts/scope_predicate_audit.py enforces this in CI. It classifies all 73 "
        "writes to profile_core and profile_engagement as directly scoped (34), "
        "scoped via a refresh-scoped TEMP source (28), or explicitly global (11), and "
        "fails on anything unaccounted for. The marker must appear within three lines "
        "above the statement.",
    )
    para(
        doc,
        "Global is sometimes correct. A deactivated NPI generates no activity, so "
        "nothing puts that person in scope; a scoped statement would never reach "
        "exactly the rows that need fixing.",
    )

    h1(doc, "5. Publication: gates and blue-green")
    para(
        doc,
        "A build does not publish by finishing. It publishes by passing.",
    )
    para(
        doc,
        "Ten named assertions run after the last step. Each is hard or soft. A hard "
        "failure sets status failed_gate, and production views continue serving the "
        "previous release untouched.",
    )
    kv_table(
        doc,
        [
            (
                "identity_source_row_delta",
                "hard",
                "Fails when the consumed identity snapshot shrank more than 10% versus the last completed build. Guards the input, not the output.",
            ),
            (
                "fill_rate_drift_critical",
                "hard",
                "Fails on a >5pp drop in a critical field's fill rate.",
            ),
            (
                "refresh_safety_check",
                "hard",
                "Proves no immutable field was mutated and no field changed without explanation.",
            ),
            (
                "orphan_satellite / population_sync",
                "hard",
                "Referential integrity between profile_core and its satellites.",
            ),
            (
                "restore_coverage",
                "hard",
                "Rebuild only. Proves app-written fields survived the rebuild.",
            ),
            (
                "anonymous_known_count_delta",
                "soft",
                "Warns on a >15% tier1 or tier2 swing.",
            ),
            ("exception_spike", "soft", "Warns on a jump in profile_exceptions."),
        ],
        headers=("Assertion", "Severity", "What it proves"),
    )
    para(
        doc,
        "Rebuild-like modes build into a candidate dataset and only repoint "
        "production views after the gates pass, so consumers never see a half-built "
        "release. profile_publish_manifest records each table promotion; "
        "profile_dataset_leases prevents two builds targeting the same dataset.",
    )
    para(
        doc,
        "An assertion is only as good as its baseline. anonymous_known_count_delta "
        "selected its baseline with mode='rebuild', and no completed rebuild has ever "
        "carried one, so it silently took the no-baseline PASS branch on every run for "
        "months -- including the day tier2 fell 40.9%. Fixed in commit ee2acce. When "
        "adding an assertion, verify the baseline query actually returns rows.",
        color=DARK_RED,
    )

    h1(doc, "6. Invariants worth knowing before editing SQL")
    kv_table(
        doc,
        [
            (
                "Step order is load-bearing",
                "populate_engagement does CREATE OR REPLACE at step 5; fill_gaps_aim_attribution does += on total_pageviews at step 12. Safe only because the table is recreated first. Moving that step into refresh would compound the addition daily.",
            ),
            (
                "Comments are -- only",
                "Pipeline SQL never uses block comments. The executor splits on statement boundaries and /* */ breaks it.",
            ),
            (
                "Snapshot and restore",
                "App-written fields are snapshotted before a rebuild and restored after, gated on their source. A rebuild without them silently empties those columns.",
            ),
            (
                "Idempotence",
                "Every enrichment should write zero rows in steady state. Prefer WHEN MATCHED AND ... IS DISTINCT FROM over unconditional UPDATE.",
            ),
            (
                "The runtime fingerprint",
                "Every run hashes its SQL files and the extractor into runtime_fingerprint. Two runs with the same fingerprint ran identical code -- useful for proving whether a box has actually pulled.",
            ),
        ],
        headers=("Invariant", "Why"),
    )

    h1(doc, "7. Telemetry, and its one blind spot")
    para(
        doc,
        "profile_build_runs holds one row per build with status, assertion summary, "
        "bytes billed, slot millis and a metadata JSON blob carrying fill rates and "
        "tier counts. profile_build_steps holds one row per step.",
    )
    para(
        doc,
        "The blind spot: a successful rebuild does not reliably appear. The "
        "2026-06-20 rebuild promoted all 16 tables in profile_publish_manifest and "
        "wrote no row to profile_build_runs at all. Only two rebuild-mode rows exist "
        "in the entire history and both are failures. Do not use profile_build_runs "
        "to answer when the last rebuild happened -- use profile_publish_manifest.",
        color=DARK_RED,
    )

    h1(doc, "8. Where to look next")
    kv_table(
        doc,
        [
            (
                "shared/profile_database_manifest.py",
                "Step and mode definitions. Start here.",
            ),
            (
                "plugins/profile_database_extractor.py",
                "The orchestrator. run_post_build_assertions() at :1964.",
            ),
            ("sql/", "One file per step. Names match the manifest."),
            (
                "profile_db_data_dictionary.docx",
                "Every table and view column. The only document covering the views.",
            ),
            ("scripts/scope_predicate_audit.py", "The CI gate on scope predicates."),
            ("docs/PROFILE_DATABASE_OPERATOR_GUIDE.md", "Operating procedure."),
        ],
        headers=("Path", "What it holds"),
    )

    EXPORTS.mkdir(parents=True, exist_ok=True)
    doc.save(PROFILE_OUT)
    print(f"[OK] Wrote {PROFILE_OUT.relative_to(REPO)}")


def main() -> int:
    build_hub()
    build_profile()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
