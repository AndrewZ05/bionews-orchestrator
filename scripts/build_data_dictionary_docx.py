"""Render the auto-generated data dictionaries as .docx for distribution.

Sources (both AUTO-GENERATED from DDL + build SQL, so they cannot go stale):
    docs/IDENTITY_HUB_DATA_DICTIONARY.md
    docs/PROFILE_DATABASE_DATA_DICTIONARY.md

Output (exports/ is gitignored, like every generated .docx):
    exports/identity_hub_data_dictionary.docx
    exports/profile_db_data_dictionary.docx

Each table gets: name, description, "Where it comes from" (the populated-by
steps from the markdown), a hand-written "How it is used" line, then the full
column table (Column | Type | Description) and any special notes.

Plain text only: no shading, no color. Bold for emphasis, monospace for
identifiers -- same conventions as build_bi_query_docx.py.

Usage:
    python scripts/build_data_dictionary_docx.py
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

REPO = Path(__file__).resolve().parents[1]
EXPORTS = REPO / "exports"

# How each table is used, in one or two sentences. Hand-written; everything
# else in the output comes from the generated markdown. A table missing here
# gets a generic line rather than a wrong one.
USAGE = {
    # identity hub
    "bn_id_persistence": "Resolve any stored bn_id before reuse: if it was retired in a merge, this returns the survivor. Every system that saves bn_ids must join through it.",
    "bn_id_xref": "The consumer surface. Audience counting, person lookup from any identifier, and the join target for the profile database and ad-platform exports. Always filter is_bot = FALSE in people counts.",
    "bn_id_hub": "The evidence log. Engineering-only: explains why identifiers were linked, including evidence too weak to act on. Not for audience numbers -- edge counts are not linkage counts.",
    "bn_id_neighbors": "Debugging surface for one person at a time: what is inside a cluster and which rule linked each pair. Always filter by bn_id; it is the largest table in the dataset.",
    "bn_id_identity_changes": "The merge/split audit trail. Explains why a person's bn_id changed and feeds graph-health monitoring (steady monthly volume is healthy).",
    "bn_id_node_index": "Pipeline machinery for incremental rebuilds: the previous run's answer key of which cluster every identifier was in. Filter to the latest promoted run or counts multiply.",
    "bn_id_metrics": "The pipeline's flight recorder: per-connector row/edge/timing stats per run. Its value is trend -- a connector's edges dropping to zero is an upstream failure surfacing here first.",
    "bn_id_manifest": "The promoted-run log. Anything reading bn_id_node_index or bn_id_metrics takes the latest promoted_at row here to know which run to trust.",
    # profile db
    "profile_core": "The master person row. Sensitive surface -- analysts should use profile_current_safe or profile_metrics instead; direct use is for approved pipelines and enrichment.",
    "profile_identifiers": "Cross-system identifier map per person, and the source of true genesis dates (MIN(first_seen) repaired created_at). Joins profiles to emails, accounts, and registry numbers.",
    "profile_engagement": "Lifetime behavioral rollup per person: email, web, forum, ad-touch. Whole-person totals, never per-site. Feeds engagement tiers and win-back lists.",
    "profile_preferences": "Newsletter subscriptions and forum settings, including the two forum clocks (registered vs last deliberate activity vs last presence).",
    "profile_survey_data": "Normalized LimeSurvey answers linked to people. Survey Engine answers live in surveyengine_data, not here.",
    "site_events": "The raw behavioral timeline (GA4-backed, since Jan 2026). Site codes are short keys; LOWER() both sides of any site join. The denominator for browsing-audience questions.",
    "profile_zero_party": "Identity-linked poll/quiz/prompt answers (largely Gravity Forms). site_domain here is the full domain, unlike site_events' short codes.",
    "profile_content_affinity": "Per-person, per-condition reading affinity. The engine behind inferred conditions (confidence 0.5) and the anonymous-opportunity analyses.",
    "profile_ad_attribution": "Ad-platform click identifiers per person (fbp/fbc/gclid family). The join surface for activation exports alongside the hub.",
    "profile_segment_tags": "Governed tags (geography, hcp verification, condition, engagement segment) promoted into interest_tags and segments.",
    "conditions_dict": "Canonical condition vocabulary (MeSH/ICD-coded). The normalization target for every condition label; aliases drive matching, so entries here are load-bearing.",
    "symptoms_dict": "HPO-coded symptom vocabulary for patient enrichment.",
    "treatments_dict": "RxNorm-coded treatment vocabulary; ready for treatment enrichment even where collection has not started.",
    "subtypes_dict": "Condition subtype vocabulary (e.g. relapsing-remitting).",
    "dictionary_meta": "Provenance and curation metadata for the dictionaries.",
    "profile_lookup": "Config-as-data: Mailchimp list-to-condition mapping, the site registry, site-to-condition mapping, newsletter catalogs. Corrected v6.7.2 after the August 2026 mapping audit.",
    "profile_build_runs": "Run-level observability: status, timings, assertion summaries. Where stuck or failed builds are diagnosed.",
    "profile_build_steps": "Step-level performance per build: duration, rows, bytes. Pinpoints which step a build died in.",
    "profile_core_snapshot": "Pre-rebuild capture of profile_core used by restore paths.",
    "profile_field_changes": "The audit log for the tracked sensitive fields: old value, new value, build, rule. How corrections stay accountable (e.g. the mapping fix's 8.3K lineage rows).",
    "profile_restore_unmapped": "Diagnostics for snapshot rows that could not be remapped to a current bn_id after a rebuild.",
    "profile_publish_manifest": "Blue/green promotion records per table per publish.",
    "profile_dataset_leases": "Mutual-exclusion leases so two rebuilds cannot publish at once.",
    "profile_evictions": "Rows removed from profile_core when identity-hub eligibility changed, with reasons.",
    "zero_party_staging": "Anonymous staging for zero-party answers before identity resolution attaches them to people.",
}

GENERIC_USAGE = (
    "Internal build surface; consult the operator guide before querying directly."
)


def parse_dictionary(md_path: Path):
    """Parse the generated markdown into (preamble, [table dicts])."""
    text = md_path.read_text(encoding="utf-8")
    # Drop the auto-generated comment header
    text = re.sub(r"<!--.*?-->\n", "", text)
    parts = re.split(r"\n## ", text)
    tables = []
    for part in parts[1:]:
        lines = part.split("\n")
        name = lines[0].strip().strip("`")
        if name in (
            "Contents",
            "1. Tables",
            "2. How rows get made: connector run order",
            "3. Columns",
        ):
            # hub dictionary has numbered sections; recurse into "3. Columns" via ###
            if name == "3. Columns":
                for sub in re.split(r"\n### ", part)[1:]:
                    tables.append(_parse_table_block(sub))
            continue
        if re.match(r"^\d+\.", name):
            continue
        tables.append(_parse_table_block(part))
    return [t for t in tables if t and t["columns"]]


def _parse_table_block(block: str):
    lines = block.split("\n")
    name = lines[0].strip().strip("`")
    desc = ""
    populated = ""
    columns = []
    notes = []
    in_notes = False
    for ln in lines[1:]:
        s = ln.strip()
        if s.startswith("**Dataset:**"):
            continue
        if s.startswith("**Description:**"):
            desc = s.replace("**Description:**", "").strip()
        elif s.startswith("**Populated by:**"):
            populated = s.replace("**Populated by:**", "").strip()
        elif s.startswith("**Columns:**"):
            continue
        elif s.startswith("**Special:**"):
            in_notes = True
        elif in_notes and s.startswith("- "):
            notes.append(s[2:])
        elif s.startswith("|") and "---" not in s:
            # definitions contain escaped pipes (enum lists); protect them
            protected = s.replace("\\|", "@@PIPE@@")
            cells = [
                c.replace("@@PIPE@@", "|").strip().strip("`")
                for c in protected.strip("|").split("|")
            ]
            if cells and cells[0] not in ("Column",):
                # hub format: Column | Type | Req | Definition ; profile: Column | Type | Definition
                if len(cells) == 4:
                    col, typ, req, definition = cells
                    if req == "yes":
                        typ = f"{typ} NOT NULL"
                    columns.append((col, typ, definition))
                elif len(cells) == 3:
                    columns.append((cells[0], cells[1], cells[2]))
        elif (
            s
            and not s.startswith("|")
            and not s.startswith("**")
            and not columns
            and not desc
        ):
            # free-text description line before the table (profile format)
            desc = s if not desc else desc
    return {
        "name": name,
        "desc": desc,
        "populated": populated,
        "columns": columns,
        "notes": notes,
    }


def _mono(run, size=8.5):
    run.font.name = "Consolas"
    run.font.size = Pt(size)


# Identity hub tables in the same order the overview and the table reference
# introduce them: the consumer surface, the evidence behind it, then the
# redirect registry and audit trail, then internals.
#
# This deliberately replaces an earlier BI-relevance ordering (ce26258). The
# four hub documents are presented as a set, in sequence, and the dictionary
# was the only one that disagreed -- bn_id_hub sat fifth here and second in
# both the overview and the table reference, so anyone reading straight
# through lost their place at the third document. Agreeing with the other
# three is worth more than being independently optimal, and the profile
# dictionary already follows its overview the same way (42fa6db).
#
# Keep this list in step with the table reference if either moves.
HUB_ORDER = [
    "bn_id_xref",
    "bn_id_hub",
    "bn_id_persistence",
    "bn_id_identity_changes",
    "bn_id_neighbors",
    "bn_id_node_index",
    "bn_id_metrics",
    "bn_id_manifest",
]


# The profile dictionary follows the story the overview tells: master row,
# its consumer surfaces, behavioral signals, self-reported answers, activation,
# pre-built segments, quality surfaces, reference data, then build internals.
# ("view", ...) entries have no physical columns here; they project from the
# tables and are described so the story reads complete.
PROFILE_STORY = [
    ("table", "profile_core"),
    ("view", "profile_current_safe",
     "The redacted person row -- one row per person with persona, condition, "
     "engagement tier, and consent, with sensitive fields (exact age, "
     "ethnicity, address, phone) redacted. The default analyst surface."),
    ("view", "profile_metrics",
     "One row per person with every headline metric precomputed as a "
     "true/false column (is_known_person, is_mailable, is_verified_hcp, "
     "is_active_email_90d, has_registration_date...). One definition per "
     "metric, so every dashboard quotes the same number. Start here."),
    ("view", "profile_marketing_audience",
     "The pre-consented audience: only people who can actually be emailed "
     "(~412K). If your query ends in a send, start here."),
    ("view", "profile_contactability",
     "What each permission allows, per person: can we email, personalize, "
     "track, advertise -- with the reason when we cannot."),
    ("table", "profile_engagement"),
    ("table", "site_events"),
    ("table", "profile_content_affinity"),
    ("table", "profile_identifiers"),
    ("table", "profile_survey_data"),
    ("table", "profile_zero_party"),
    ("table", "profile_preferences"),
    ("table", "profile_ad_attribution"),
    ("table", "profile_segment_tags"),
    ("view", "profile_audience_hcp",
     "Pre-built segment: verified clinicians (~362K). Use it instead of "
     "rebuilding the segment logic."),
    ("view", "profile_audience_patients_confirmed",
     "Pre-built segment: confirmed patients (~77K) -- people with a "
     "deterministic or self-declared patient signal, not inferred readers."),
    ("view", "profile_audience_caregivers",
     "Pre-built segment: caregivers (~10.6K)."),
    ("view", "profile_audience_high_engagement",
     "Pre-built segment: the high engagement tier (top ~5 percent)."),
    ("view", "profile_coverage",
     "Quality surface: how filled each field is, per scope -- the population "
     "the check runs over (all, engagement, hcp, patient). Check before "
     "building on a field you have not used."),
    ("view", "profile_exceptions",
     "Quality surface: classification conflicts the build logged, one row per "
     "affected person per exception type."),
    ("table", "conditions_dict"),
    ("table", "profile_lookup"),
    ("table", "symptoms_dict"),
    ("table", "treatments_dict"),
    ("table", "subtypes_dict"),
    ("table", "dictionary_meta"),
    ("table", "profile_field_changes"),
    ("table", "profile_build_runs"),
    ("table", "profile_build_steps"),
    ("table", "profile_publish_manifest"),
    ("table", "profile_core_snapshot"),
    ("table", "profile_evictions"),
    ("table", "profile_restore_unmapped"),
    ("table", "profile_dataset_leases"),
    ("table", "zero_party_staging"),
]


VIEW_PROJECT = "bi-data-391216"
VIEW_DATASET = "profile_data"


def fetch_view_columns(view_names):
    """Live column list per view from BigQuery INFORMATION_SCHEMA, including
    the column descriptions applied by apply_profile_metadata. Returns {} when
    credentials are unavailable so generation still succeeds (views then render
    without column tables)."""
    try:
        import sys

        sys.path.insert(0, str(REPO))
        from shared.bigquery_client import get_bigquery_client, setup_gcp_credentials

        setup_gcp_credentials()
        client = get_bigquery_client(project=VIEW_PROJECT)
        names = ", ".join(f"'{n}'" for n in view_names)
        rows = client.query(
            f"""
            SELECT c.table_name, c.column_name, c.data_type, c.ordinal_position,
                   fp.description
            FROM `{VIEW_PROJECT}.{VIEW_DATASET}.INFORMATION_SCHEMA.COLUMNS` c
            LEFT JOIN `{VIEW_PROJECT}.{VIEW_DATASET}.INFORMATION_SCHEMA.COLUMN_FIELD_PATHS` fp
              ON fp.table_name = c.table_name AND fp.field_path = c.column_name
            WHERE c.table_name IN ({names})
            ORDER BY c.table_name, c.ordinal_position
            """
        ).result()
        out = {}
        for r in rows:
            typ = r.data_type
            if len(typ) > 60:
                typ = typ[:57] + "..."
            out.setdefault(r.table_name, []).append(
                (r.column_name, typ, r.description or "")
            )
        return out
    except Exception as exc:  # noqa: BLE001
        print(f"  [WARN] view columns unavailable ({exc.__class__.__name__}); "
              "views will render without column tables")
        return {}


def build(md_path: Path, out_path: Path, title: str, intro: str, order=None, story=None) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    tables = parse_dictionary(md_path)
    if story:
        by_name = {t["name"]: t for t in tables}
        ordered = []
        for item in story:
            if item[0] == "table" and item[1] in by_name:
                ordered.append(by_name.pop(item[1]))
            elif item[0] == "view":
                ordered.append({"name": item[1], "desc": item[2], "populated": "",
                                "columns": [], "notes": [], "is_view": True})
        ordered.extend(by_name.values())
        tables = ordered
    elif order:
        rank = {name: i for i, name in enumerate(order)}
        tables.sort(key=lambda t: rank.get(t["name"], len(order)))
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph(intro)
    doc.add_paragraph(
        "For each table: what it holds, where its rows come from (the build "
        "steps or connectors that write it), how it is used, then every column "
        "with its type and definition. Column definitions frequently name the "
        "originating source system (NPI, Mailchimp, LimeSurvey, GA4, BuddyPress) "
        "-- that is the per-column provenance."
    )

    view_names = [t["name"] for t in tables if t.get("is_view")]
    view_columns = fetch_view_columns(view_names) if view_names else {}

    for t in tables:
        kind = "View" if t.get("is_view") else "Table"
        doc.add_heading(f"{kind}: {t['name']}", level=1)
        if t.get("is_view"):
            doc.add_paragraph(t["desc"])
            p = doc.add_paragraph()
            r = p.add_run(
                "View -- projects from the tables in this dictionary; the "
                "column list below is read live from the warehouse at "
                "generation time."
            )
            r.italic = True
            r.font.size = Pt(9)
            vcols = view_columns.get(t["name"], [])
            if vcols:
                t = dict(t)
                t["columns"] = vcols
            else:
                continue
        if t["desc"] and not t.get("is_view"):
            doc.add_paragraph(t["desc"])
        if t["populated"]:
            p = doc.add_paragraph()
            p.add_run("Where it comes from: ").bold = True
            r = p.add_run(t["populated"])
            r.font.size = Pt(9)
        p = doc.add_paragraph()
        p.add_run("How it is used: ").bold = True
        p.add_run(USAGE.get(t["name"], GENERIC_USAGE))

        dt = doc.add_table(rows=1, cols=3)
        dt.style = "Table Grid"
        for i, h in enumerate(("Column", "Type", "Description")):
            cell = dt.cell(0, i)
            cell.text = ""
            r = cell.paragraphs[0].add_run(h)
            r.bold = True
            r.font.size = Pt(9)
        for col, typ, definition in t["columns"]:
            row = dt.add_row()
            row.cells[0].text = ""
            _mono(row.cells[0].paragraphs[0].add_run(col.replace("&nbsp;", " ")))
            row.cells[1].text = ""
            _mono(row.cells[1].paragraphs[0].add_run(typ))
            row.cells[2].text = ""
            r = row.cells[2].paragraphs[0].add_run(definition)
            r.font.size = Pt(9)
        for note in t["notes"]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(note.replace("**", ""))
            r.font.size = Pt(9)

    doc.save(out_path)
    print(f"[OK] Wrote {out_path.relative_to(REPO)} ({len(tables)} tables)")


def main() -> int:
    build(
        REPO / "docs" / "IDENTITY_HUB_DATA_DICTIONARY.md",
        EXPORTS / "identity_hub_data_dictionary.docx",
        "Identity Hub Data Dictionary",
        "Every table in identity_hub_data, in the order the overview and the "
        "table reference introduce them: bn_id_xref first because it covers "
        "most BI needs on its own, then bn_id_hub as the evidence behind it, "
        "then the redirect registry and audit trail any report might touch, "
        "then engineering-only internals. "
        "Rendered from the auto-generated markdown dictionary, which is "
        "produced from the DDL, the config, and the pipeline code itself -- "
        "regenerate both with scripts/generate_identity_hub_docs.py then this "
        "script.",
        order=HUB_ORDER,
    )
    build(
        REPO / "docs" / "PROFILE_DATABASE_DATA_DICTIONARY.md",
        EXPORTS / "profile_db_data_dictionary.docx",
        "Profile Database Data Dictionary",
        "Every physical table in profile_data, profile_ops and profile_staging, "
        "ordered by relevance to the BI team: the person row and behavioral "
        "signal tables first, then self-reported answers and preferences, then "
        "activation surfaces and reference data, then build internals. "
        "Rendered from the auto-generated markdown dictionary (v6.7), which is "
        "produced from the DDL and the active SQL WRITES headers -- regenerate "
        "both with scripts/generate_schema_contract.py then this script. "
        "Views (profile_current_safe, profile_metrics, the audience views) "
        "project from these tables; see the table_reference and bi_queries "
        "documents for the view directory.",
        story=PROFILE_STORY,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
