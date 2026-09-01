#!/usr/bin/env python3
"""
Build exports/platform_two_week_summary.docx.

What was delivered to the Identity Hub and Profile Database between
2026-08-10 and 2026-08-24, written as feature then benefit so a non-technical
reader can see what each change actually gets us.

American English throughout. Every figure production-verified 2026-08-24.

  python scripts/build_two_week_summary_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
EXPORTS = REPO / "exports"
OUT = EXPORTS / "platform_two_week_summary.docx"

NAVY = RGBColor(0x0D, 0x47, 0xA1)
GREY = RGBColor(0x5F, 0x63, 0x68)


def para(doc, text, size=11, bold=False, italic=False, color=None, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = NAVY
        r.font.size = Pt(14)
    return h


def labeled(doc, label, text):
    """Feature/benefit paragraph: bold lead-in, then the content."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(label + "  ")
    r.font.size = Pt(11)
    r.font.bold = True
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p


SECTIONS = [
    (
        "We can now tell you when someone joined us",
        "Acquisition dates for nearly every person in the database, recovered from "
        "identity records, cookie creation dates and several fallback sources. Coverage "
        "went from 64.6 percent to 99.6 percent of known people. We then purged dates "
        "that were provably impossible, so what remains is trustworthy rather than "
        "merely present.",
        "We can report growth. Before this we could tell you how many people we had but "
        "not when they arrived, which meant no month over month reporting, no way to "
        "measure whether a campaign brought anyone in, and no way to tell genuine new "
        "audience from the same people returning. That is now a routine question. We "
        "also built growth views with the right filters applied by default, so the "
        "number is correct without anyone remembering to exclude the wrong things.",
        [],
    ),
    (
        "We remember people for thirteen months instead of ninety days",
        "Extended how long the identity system retains someone who has gone quiet, from "
        "an effective ninety days to thirteen months.",
        "Our audience does not behave like a retail audience. Someone researches a "
        "diagnosis intensively for a month, goes quiet while they get on with treatment, "
        "then comes back when something changes. Under the old window we forgot them and "
        "treated the return as a brand new stranger, losing their condition, their "
        "history and their preferences. 4.5 million of the people we recognize today "
        "were last seen more than ninety days ago, and 529,890 were last seen more than "
        "a year ago. All of them would have been strangers.",
        [],
    ),
    (
        "Every Mailchimp subscriber is now connected to a person",
        "Changed how Mailchimp members are anchored so that every subscriber is "
        "represented in the identity graph, whether or not we have seen them browse.",
        "Subscribers missing from the graph fell from around 3,529 to 83, a 98 percent "
        "recovery. A subscriber who is missing cannot be matched to their profile, which "
        "means they are invisible to audience counts and cannot be personalized to. "
        "These are engaged people who open our email, so they are exactly the audience "
        "we least want to lose track of.",
        [],
    ),
    (
        "Our clinician numbers are now honest",
        "Separated the healthcare professional audience into measures that answer "
        "different questions, and stopped counting clinicians whose credential has "
        "since lapsed.",
        "One number used to do every job, and it was the largest one. Quoting it "
        "commercially overstated our reachable clinician audience roughly sixteenfold. "
        "The system now makes the right number the easy one to reach for, which is the "
        "discipline you want when the answer is going into a client conversation. We "
        "also hold specialty for 322,882 of them and practice location for 326,040, so "
        "the audience can be segmented meaningfully rather than treated as one block.",
        [
            "Verified clinicians: 380,508",
            "Verified against the federal provider registry: 356,859",
            "Who have engaged with our content: 102,584",
            "Who we can reach by email: 23,027",
            "Active on email in the last ninety days: 13,935",
        ],
    ),
    (
        "We fixed a large amount of incorrect data",
        "A series of corrections across the database.",
        "Every one of these was quietly making a number wrong somewhere. The condition "
        "fix alone returned 277,887 people to a condition audience they belonged in. The "
        "consent fix returned 5,220 opted in people to our reachable count. These are "
        "people we already had and were failing to see.",
        [
            "Corrected 277,887 stranded conditions, people whose condition had been "
            "recorded and then lost.",
            "Fixed 1,540 stale provider deactivation records, then put all eight "
            "provider data columns onto daily maintenance so they cannot silently go "
            "stale again.",
            "Corrected the mapping between Mailchimp lists and conditions.",
            "Moved engagement scoring onto bot filtered opens rather than raw opens.",
            "Stopped automated email prefetches from making dormant people look active.",
            "Corrected an email consent issue where a cookie preference was wrongly "
            "blocking 5,220 people who had explicitly opted in.",
            "Corrected forum activity, which had been reading the wrong source and "
            "undercounting badly.",
            "Replaced an empty field with content affinity, which actually covers "
            "59,655 clinicians.",
            "Made dictionary drift self healing rather than silently diverging.",
        ],
    ),
    (
        "We connected our own registration platform",
        "A complete data pipeline for SurveyEngine, seventeen tables, wired into both "
        "the identity system and the profile database. Registration answers now enrich "
        "profiles directly.",
        "SurveyEngine is becoming the primary way people register with us. When somebody "
        "registers directly we learn far more, and we learn it because they told us "
        "rather than because we inferred it. As it becomes the main route in, the "
        "quality of what we know rises across the whole audience. The connection is "
        "built and waiting, so that improvement begins the moment volume shifts.",
        [],
    ),
    (
        "We agreed one definition per measure",
        "A single place where every audience measure is defined once. Known person, "
        "verified clinician, engaged clinician, mailable, active on email, forum member, "
        "and others.",
        "Two people asking the same question now get the same answer. Before this, every "
        "analyst rebuilt definitions in their own query and produced slightly different "
        "numbers, and there was no way to tell which was right. Definitions worth "
        "arguing about now live in one place where they can be argued about once.",
        [],
    ),
    (
        "We made the system catch its own mistakes",
        "A set of safeguards around the daily build.",
        "The system refuses to publish numbers it cannot vouch for. When something looks "
        "wrong, production keeps serving the last good data while somebody investigates, "
        "rather than quietly serving something worse. The snapshots proved their worth "
        "within days of being introduced, when they became the recovery path for an "
        "identity issue.",
        [
            "A gate on the identity input, so a problem upstream is caught before it "
            "reaches anyone rather than after.",
            "A fix to a safety check that had been silently passing for months because "
            "it was comparing against a baseline that did not exist.",
            "Monthly forensic snapshots of both systems, with verification and alerting.",
            "A tightened safety threshold, from 0.5 to 0.9, on how much a table may "
            "shrink before a write is refused.",
            "Resilience to large population jumps, so a big upstream change does not "
            "fail the daily run.",
            "A fix to a filter that was deleting heavily linked people instead of "
            "trimming their weakest connections.",
            "Silent routing failures now fail loudly rather than reporting success.",
        ],
    ),
]

DOCS = [
    (
        "For the BI team",
        [
            "A plain English overview of each system.",
            "A table reference for each, grouped by what the tables are for.",
            "A complete data dictionary covering every column and every view.",
            "A query guide with fifteen worked queries for profiles and fourteen for "
            "identity. Each states the mistake it prevents before it shows the query.",
            "Both query sets also as runnable files, so nobody has to copy anything out "
            "of a document.",
        ],
    ),
    (
        "For engineering",
        [
            "Reference documents for both systems covering how they are actually built, "
            "the algorithms, the failure modes seen in production, and the invariants "
            "that are easy to break.",
            "A presentation deck covering the same ground.",
        ],
    ),
    (
        "For executives and partners",
        [
            "A presentation deck telling the story of both systems as one platform.",
            "An executive overview document.",
        ],
    ),
    (
        "For operators",
        [
            "Runbooks covering the full identity rebuild, what to check before running "
            "one, what the stop conditions are, and how to recover.",
            "Incident records covering what happened, what caused it, and what rule came "
            "out of it.",
        ],
    ),
]


def build():
    doc = Document()

    doc.add_heading("Identity Hub and Profile Database", level=0)
    para(
        doc,
        "What we delivered, August 10 to 24, 2026",
        size=13,
        italic=True,
        color=GREY,
    )
    para(doc, "Executive summary", size=10, color=GREY, after=16)

    para(
        doc,
        "Here is what went in over the past two weeks, written as what we built and "
        "what it gets us. 115 changes in total.",
    )

    for title, feature, benefit, details in SECTIONS:
        h1(doc, title)
        labeled(doc, "What we built:", feature)
        if details:
            for d in details:
                bullet(doc, d)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        labeled(doc, "What it gets us:", benefit)

    h1(doc, "We documented all of it")
    labeled(
        doc,
        "What we built:",
        "Ten reference documents, two presentation decks, two sets of runnable "
        "queries and a set of operator runbooks, written for four different "
        "audiences.",
    )
    for aud, items in DOCS:
        para(doc, aud, bold=True, after=3)
        for i in items:
            bullet(doc, i)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
    labeled(
        doc,
        "What it gets us:",
        "The BI team can answer their own questions without waiting on anyone, and "
        "without accidentally quoting a number that overstates what we have. A new "
        "engineer can understand the system without reverse engineering it. And the "
        "knowledge is not in one person's head.",
    )

    h1(doc, "We cleaned up")
    labeled(
        doc,
        "What we built:",
        "Retired unused connectors and removed 2.35 million dead links, dropped "
        "legacy columns, removed superseded views, untracked 67 out of date "
        "documents, and disabled a source that is on hold.",
    )
    labeled(
        doc,
        "What it gets us:",
        "Less to maintain, less to misread and less that can silently break. Every "
        "retired connector was verified to have no effect on identity before "
        "removal.",
    )

    h1(doc, "One honest note")
    para(
        doc,
        "The August 21 identity rebuild surfaced two defects. Both are fixed, and both "
        "were caught by our own safeguards before anything reached production. No data "
        "was lost, and no consent or medical information was affected. Recovery was "
        "possible because of the snapshots we had introduced the day before.",
    )
    para(doc, "Happy to walk anyone through any part of this in more detail.")

    EXPORTS.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"[OK] Wrote {OUT.relative_to(REPO)}  ({words:,} words)")
    return 0


if __name__ == "__main__":
    raise SystemExit(build())
