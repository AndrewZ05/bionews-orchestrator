"""Generate Threads Pipeline Specification DOCX."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import date
import os

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_header_row(table, headers):
    """Format first row of table as header with dark blue background."""
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, '2B579A')

def add_data_row(table, row_idx, values, bold_first=False):
    """Add data to a table row."""
    row = table.rows[row_idx]
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(v))
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.bold = True
        # Alternate row shading
        if row_idx % 2 == 0:
            set_cell_shading(cell, 'F2F2F2')

def set_table_style(table):
    """Apply consistent table formatting."""
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)

def build_doc():
    doc = Document()

    # ── Default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    # ═══════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ═══════════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Threads API ETL Pipeline')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Technical Specification')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph('')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'Version 1.0  |  {date.today().strftime("%B %d, %Y")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS (manual)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1', 'Overview'),
        ('2', 'Architecture'),
        ('3', 'API Integration'),
        ('4', 'Authentication & Token Management'),
        ('5', 'Data Model'),
        ('6', 'Extraction Logic'),
        ('7', 'Data Transformations (SQL)'),
        ('8', 'Rate Limiting & Retry Strategy'),
        ('9', 'Multi-Account Architecture'),
        ('10', 'Configuration Reference'),
        ('11', 'Scheduling & Operations'),
        ('12', 'Error Handling'),
        ('13', 'Testing & App Review'),
        ('14', 'Known Limitations & Gotchas'),
        ('15', 'Appendix: Environment Variables'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}.  {title_text}')
        run.font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 1. OVERVIEW
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('1. Overview', level=1)

    doc.add_paragraph(
        'This document specifies the Threads API ETL pipeline, a component of the Orchestrator '
        'platform that extracts social media data from Meta\'s Threads API, transforms it, and '
        'loads it into Google BigQuery for analytics and reporting.'
    )

    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        'The pipeline provides automated, incremental extraction of Threads content and engagement '
        'metrics across multiple Threads accounts. It supports daily scheduling, full and incremental '
        'refresh modes, and produces a unified analytics layer in BigQuery.'
    )

    doc.add_heading('1.2 Scope', level=2)
    items = [
        'Profile metadata for all configured Threads accounts',
        'Posts (threads) with full media metadata',
        'Per-post engagement insights (views, likes, replies, reposts, quotes, shares)',
        'Daily account-level insights (views, likes, replies, reposts, quotes, followers_count)',
        'Derived analytics: engagement rate, total engagements',
        'Multi-account token lifecycle management (OAuth, refresh, rotation)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('1.3 Key Differentiators from Facebook/Instagram', level=2)
    doc.add_paragraph(
        'The Threads API is a completely separate API from the Facebook Graph API. It uses a '
        'different domain (graph.threads.net), different OAuth grant types (th_exchange_token, '
        'th_refresh_token), and is not supported by the facebook_business Python SDK. The pipeline '
        'uses raw HTTP requests via the requests library.'
    )

    t = doc.add_table(rows=8, cols=3)
    add_header_row(t, ['Aspect', 'Facebook / Instagram', 'Threads'])
    data = [
        ('API Domain', 'graph.facebook.com', 'graph.threads.net/v1.0'),
        ('SDK Support', 'facebook_business SDK', 'None (raw requests)'),
        ('OAuth Grant', 'fb_exchange_token', 'th_exchange_token'),
        ('Token Refresh', 'refresh_token', 'th_refresh_token'),
        ('Account Model', 'Pages + Ad Accounts', 'Per-user (no Pages)'),
        ('Media Types', 'PHOTO, VIDEO, CAROUSEL, etc.', 'TEXT, IMAGE, VIDEO, CAROUSEL_ALBUM, REPOST_FACADE'),
        ('Insights', 'Breakdown by paid/organic/fan', 'No breakdowns; REPOST_FACADE excluded'),
    ]
    for i, row_data in enumerate(data):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 2. ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('2. Architecture', level=1)

    doc.add_heading('2.1 Component Overview', level=2)
    doc.add_paragraph(
        'The Threads pipeline follows the Orchestrator\'s plugin-based architecture. '
        'It consists of four primary components:'
    )

    t = doc.add_table(rows=5, cols=3)
    add_header_row(t, ['Component', 'File', 'Responsibility'])
    data = [
        ('Extractor Plugin', 'plugins/threads_extractor.py', 'ETL orchestration, data extraction, Parquet export'),
        ('API Client', 'shared/threads_client.py', 'HTTP session, rate limiting, pagination, retries'),
        ('Token Manager', 'shared/threads_token_manager.py', 'OAuth flow, token refresh, multi-account store'),
        ('Configuration', 'configs/threads.yaml', 'Schema definitions, pipeline settings, scheduling'),
    ]
    for i, row_data in enumerate(data):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_heading('2.2 Data Flow', level=2)
    doc.add_paragraph(
        'The end-to-end data flow follows the standard Orchestrator pattern:'
    )

    steps = [
        'Token validation & auto-refresh — verify all account tokens before extraction begins',
        'API extraction — call Threads API endpoints for each account and table',
        'Record normalization — parse timestamps, stringify IDs, add execution metadata',
        'Parquet export — write normalized records to local Parquet files',
        'GCS upload — upload Parquet files to Google Cloud Storage',
        'BigQuery staging — load Parquet into staging tables',
        'Hash merge — deduplicate and merge staging into production tables',
        'Post-process SQL — execute threads_post_insights_daily.sql to build analytics views',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}', style='List Number')

    doc.add_heading('2.3 Plugin Loading', level=2)
    doc.add_paragraph(
        'The orchestrator dynamically loads the Threads extractor via:\n'
        '    importlib.import_module("plugins.threads_extractor")\n\n'
        'The source type "threads" is registered in shared/config_loader.py (line 237). '
        'The plugin exposes a run_pipeline() function with the standard Orchestrator signature.'
    )

    doc.add_heading('2.4 Table Dependencies', level=2)
    doc.add_paragraph(
        'Tables have extraction dependencies that determine processing order:'
    )
    deps = [
        ('profiles', 'None — extracted first'),
        ('posts', 'Depends on profiles (needs threads_user_id)'),
        ('post_insights', 'Depends on posts (needs thread_id from each post)'),
        ('user_insights', 'Depends on profiles (enriches with username)'),
    ]
    t = doc.add_table(rows=5, cols=2)
    add_header_row(t, ['Table', 'Depends On'])
    for i, row_data in enumerate(deps):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 3. API INTEGRATION
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('3. API Integration', level=1)

    doc.add_heading('3.1 Base URL', level=2)
    doc.add_paragraph('All API calls target:  https://graph.threads.net/v1.0')

    doc.add_heading('3.2 Endpoints', level=2)

    t = doc.add_table(rows=8, cols=4)
    add_header_row(t, ['Endpoint', 'Method', 'Purpose', 'Key Parameters'])
    endpoints = [
        ('/me', 'GET', 'Authenticated user profile', 'fields=id,username,name,threads_profile_picture_url,threads_biography'),
        ('/{user_id}/threads', 'GET', 'List user posts', 'fields, since, until, limit=100'),
        ('/{thread_id}/insights', 'GET', 'Per-post engagement', 'metric=views,likes,replies,reposts,quotes,shares'),
        ('/{user_id}/threads_insights', 'GET', 'Daily account metrics', 'metric, since, until, period=day'),
        ('/{thread_id}/replies', 'GET', 'Post replies', 'fields, limit'),
        ('/{thread_id}/conversation', 'GET', 'Nested conversation', 'fields, limit'),
        ('/refresh_access_token', 'GET', 'Refresh long-lived token', 'grant_type=th_refresh_token, access_token'),
    ]
    for i, row_data in enumerate(endpoints):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('3.3 OAuth Endpoints', level=2)

    t = doc.add_table(rows=4, cols=4)
    add_header_row(t, ['Endpoint', 'Method', 'Purpose', 'Key Parameters'])
    oauth_eps = [
        ('threads.net/oauth/authorize', 'Browser', 'User authorization', 'client_id, redirect_uri, scope, response_type=code'),
        ('graph.threads.net/oauth/access_token', 'POST', 'Exchange code for short-lived token', 'client_id, client_secret, grant_type=authorization_code, code'),
        ('graph.threads.net/access_token', 'GET', 'Exchange for long-lived token', 'grant_type=th_exchange_token, client_secret, access_token'),
    ]
    for i, row_data in enumerate(oauth_eps):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('3.4 Required Permissions (Scopes)', level=2)
    perms = [
        ('threads_basic', 'Profile information, list media/posts'),
        ('threads_manage_insights', 'Post-level and account-level engagement insights'),
        ('threads_read_replies', 'Replies and conversation threads'),
    ]
    t = doc.add_table(rows=4, cols=2)
    add_header_row(t, ['Scope', 'Grants Access To'])
    for i, row_data in enumerate(perms):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_heading('3.5 Pagination', level=2)
    doc.add_paragraph(
        'The Threads API uses Meta\'s standard cursor-based pagination. Each response includes '
        'a paging.next URL that the client follows until exhausted. The client library '
        '(paginate_threads_api) handles this automatically, yielding records one at a time. '
        'Page size is fixed at 100 items per request.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 4. AUTHENTICATION & TOKEN MANAGEMENT
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('4. Authentication & Token Management', level=1)

    doc.add_heading('4.1 Token Lifecycle', level=2)
    stages = [
        'Authorization — User grants access via OAuth browser flow or User Token Generator',
        'Short-lived token — Valid for 1 hour; exchanged immediately for long-lived token',
        'Long-lived token — Valid for 60 days; auto-refreshed when within 7 days of expiry',
        'Refresh — GET /refresh_access_token with grant_type=th_refresh_token extends for another 60 days',
        'Expiration — If not refreshed within 60 days, token becomes invalid and re-authorization is required',
    ]
    for i, s in enumerate(stages, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')

    doc.add_heading('4.2 Multi-Account Token Store', level=2)
    doc.add_paragraph(
        'Tokens for all Threads accounts are stored in .threads_tokens.json at the project root. '
        'Each entry is keyed by username and contains:'
    )
    fields = [
        ('user_id', 'Threads user ID (numeric string)'),
        ('access_token', 'Long-lived OAuth token (prefix THAA...)'),
        ('expires_at', 'Unix timestamp of token expiration'),
        ('created_at', 'ISO 8601 timestamp of when token was added/refreshed'),
    ]
    t = doc.add_table(rows=5, cols=2)
    add_header_row(t, ['Field', 'Description'])
    for i, row_data in enumerate(fields):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_heading('4.3 Legacy Single-Account Fallback', level=2)
    doc.add_paragraph(
        'If .threads_tokens.json is absent or empty, the extractor falls back to environment '
        'variables THREADS_ACCESS_TOKEN and THREADS_USER_ID from the .env file. Token metadata '
        '(expiry tracking) is stored separately in .threads_token_meta.json.'
    )

    doc.add_heading('4.4 Auto-Refresh Logic', level=2)
    doc.add_paragraph(
        'Before each extraction run, the pipeline checks every account token:\n\n'
        '  - If remaining days > 7: no action (token is healthy)\n'
        '  - If remaining days <= 7 and > 0: auto-refresh via API, save new token to store\n'
        '  - If remaining days <= 0: log warning (token expired), skip account\n\n'
        'Auto-refresh is triggered per-account at the start of run_pipeline().'
    )

    doc.add_heading('4.5 Token Manager CLI', level=2)
    cmds = [
        ('--authorize', 'Start full OAuth flow (opens browser)'),
        ('--add-token TOKEN', 'Add a token from Meta User Token Generator'),
        ('--refresh', 'Refresh all account tokens (or --account NAME for one)'),
        ('--refresh --force', 'Force refresh even if token has > 7 days remaining'),
        ('--status', 'Display all account token statuses with live API check'),
        ('--status --account NAME', 'Display status for a specific account'),
    ]
    t = doc.add_table(rows=7, cols=2)
    add_header_row(t, ['Command', 'Description'])
    for i, row_data in enumerate(cmds):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_paragraph(
        '\nUsage:  python shared/threads_token_manager.py <command>'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 5. DATA MODEL
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('5. Data Model', level=1)

    doc.add_heading('5.1 BigQuery Datasets', level=2)
    datasets = [
        ('threads_staging', 'Staging tables — raw data loaded from Parquet before hash merge'),
        ('threads_data', 'Production tables — deduplicated data after hash merge'),
        ('threads_test', 'Test tables — used when running in test mode'),
        ('BN_Warehouse', 'Analytics layer — post-process SQL creates derived tables here'),
    ]
    t = doc.add_table(rows=5, cols=2)
    add_header_row(t, ['Dataset', 'Purpose'])
    for i, row_data in enumerate(datasets):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_heading('5.2 Table: threads_profiles', level=2)
    doc.add_paragraph('Profile metadata for each Threads account. Refreshed on every run.')
    cols = [
        ('threads_user_id', 'STRING', 'PK', 'Threads numeric user ID'),
        ('username', 'STRING', '', 'Threads handle (e.g., @msnewstoday)'),
        ('name', 'STRING', '', 'Display name'),
        ('threads_profile_picture_url', 'STRING', '', 'Avatar URL'),
        ('threads_biography', 'STRING', '', 'Profile bio text'),
        ('source', 'STRING', '', 'Pipeline source identifier'),
        ('execution_id', 'STRING', '', 'Unique pipeline run ID'),
        ('extracted_at', 'TIMESTAMP', '', 'Extraction timestamp'),
    ]
    t = doc.add_table(rows=len(cols) + 1, cols=4)
    add_header_row(t, ['Column', 'Type', 'Key', 'Description'])
    for i, row_data in enumerate(cols):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('5.3 Table: threads_posts', level=2)
    doc.add_paragraph('Individual threads (posts) with full media metadata.')
    cols = [
        ('threads_user_id', 'STRING', 'PK', 'Author Threads user ID'),
        ('id', 'STRING', 'PK', 'Thread/post ID'),
        ('media_product_type', 'STRING', '', 'Always "THREADS"'),
        ('media_type', 'STRING', '', 'TEXT | IMAGE | VIDEO | CAROUSEL_ALBUM | REPOST_FACADE'),
        ('media_url', 'STRING', '', 'Media asset URL (images/videos)'),
        ('permalink', 'STRING', '', 'Threads.net URL to the post'),
        ('text', 'STRING', '', 'Post text content'),
        ('timestamp', 'TIMESTAMP', '', 'Post creation time (ISO 8601)'),
        ('shortcode', 'STRING', '', 'Short identifier for the post'),
        ('thumbnail_url', 'STRING', '', 'Video thumbnail URL'),
        ('is_quote_post', 'STRING', '', 'Whether this is a quote post'),
        ('alt_text', 'STRING', '', 'Alt text for media'),
        ('link_attachment_url', 'STRING', '', 'Attached link URL'),
        ('has_replies', 'STRING', '', 'Whether post has replies'),
        ('is_reply', 'STRING', '', 'Whether this is a reply to another thread'),
        ('source', 'STRING', '', 'Pipeline source identifier'),
        ('execution_id', 'STRING', '', 'Unique pipeline run ID'),
        ('extracted_at', 'TIMESTAMP', '', 'Extraction timestamp'),
    ]
    t = doc.add_table(rows=len(cols) + 1, cols=4)
    add_header_row(t, ['Column', 'Type', 'Key', 'Description'])
    for i, row_data in enumerate(cols):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('5.4 Table: threads_post_insights', level=2)
    doc.add_paragraph(
        'Engagement metrics per post. REPOST_FACADE media type is excluded (API does not support insights for reposts).'
    )
    cols = [
        ('threads_user_id', 'STRING', 'PK', 'Author Threads user ID'),
        ('thread_id', 'STRING', 'PK', 'Thread/post ID'),
        ('timestamp', 'TIMESTAMP', '', 'Post creation timestamp'),
        ('views', 'STRING', '', 'Number of times the post was viewed'),
        ('likes', 'STRING', '', 'Number of likes'),
        ('replies', 'STRING', '', 'Number of replies'),
        ('reposts', 'STRING', '', 'Number of reposts'),
        ('quotes', 'STRING', '', 'Number of quote posts'),
        ('shares', 'STRING', '', 'Number of shares'),
        ('source', 'STRING', '', 'Pipeline source identifier'),
        ('execution_id', 'STRING', '', 'Unique pipeline run ID'),
        ('extracted_at', 'TIMESTAMP', '', 'Extraction timestamp'),
    ]
    t = doc.add_table(rows=len(cols) + 1, cols=4)
    add_header_row(t, ['Column', 'Type', 'Key', 'Description'])
    for i, row_data in enumerate(cols):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('5.5 Table: threads_user_insights', level=2)
    doc.add_paragraph(
        'Daily account-level engagement metrics. Date range is chunked into 90-day windows. '
        'followers_count requires the account to have 100+ followers.'
    )
    cols = [
        ('threads_user_id', 'STRING', 'PK', 'Threads user ID'),
        ('date', 'TIMESTAMP', 'PK', 'Metric date (daily grain)'),
        ('username', 'STRING', '', 'Threads handle (enriched from profiles)'),
        ('views', 'STRING', '', 'Daily view count'),
        ('likes', 'STRING', '', 'Daily like count'),
        ('replies', 'STRING', '', 'Daily reply count'),
        ('reposts', 'STRING', '', 'Daily repost count'),
        ('quotes', 'STRING', '', 'Daily quote count'),
        ('followers_count', 'STRING', '', 'Follower count (requires 100+ followers)'),
        ('source', 'STRING', '', 'Pipeline source identifier'),
        ('execution_id', 'STRING', '', 'Unique pipeline run ID'),
        ('extracted_at', 'TIMESTAMP', '', 'Extraction timestamp'),
    ]
    t = doc.add_table(rows=len(cols) + 1, cols=4)
    add_header_row(t, ['Column', 'Type', 'Key', 'Description'])
    for i, row_data in enumerate(cols):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('5.6 Derived Table: threads_post_insights_daily', level=2)
    doc.add_paragraph(
        'Analytics fact table in BN_Warehouse. Created by the post-process SQL step. '
        'Joins posts with their insights and computes derived engagement metrics.'
    )
    doc.add_paragraph('Partitioned by: date')
    doc.add_paragraph('Clustered by: threads_user_id, thread_id')
    cols = [
        ('threads_user_id', 'STRING', 'Author'),
        ('thread_id', 'STRING', 'Post ID (from posts.id)'),
        ('text', 'STRING', 'Post content'),
        ('media_type', 'STRING', 'TEXT | IMAGE | VIDEO | CAROUSEL_ALBUM | REPOST_FACADE'),
        ('permalink', 'STRING', 'Post URL'),
        ('timestamp', 'TIMESTAMP', 'Post creation time'),
        ('date', 'DATE', 'Partition key (derived from timestamp)'),
        ('views', 'FLOAT64', 'View count (0 if NULL)'),
        ('likes', 'FLOAT64', 'Like count'),
        ('replies', 'FLOAT64', 'Reply count'),
        ('reposts', 'FLOAT64', 'Repost count'),
        ('quotes', 'FLOAT64', 'Quote count'),
        ('shares', 'FLOAT64', 'Share count'),
        ('total_engagements', 'FLOAT64', 'likes + replies + reposts + quotes + shares'),
        ('engagement_rate', 'FLOAT64', 'total_engagements / views (SAFE_DIVIDE)'),
    ]
    t = doc.add_table(rows=len(cols) + 1, cols=3)
    add_header_row(t, ['Column', 'Type', 'Description'])
    for i, row_data in enumerate(cols):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('5.7 Hash Merge Keys', level=2)
    doc.add_paragraph(
        'The hash merge process deduplicates records using composite primary keys. '
        'The following fields are excluded from the hash computation: loaded_at, extracted_at, '
        'source, execution_id, updated_at, row_hash.'
    )
    keys = [
        ('threads_profiles', '[threads_user_id]'),
        ('threads_posts', '[threads_user_id, id]'),
        ('threads_post_insights', '[threads_user_id, thread_id]'),
        ('threads_user_insights', '[threads_user_id, date]'),
    ]
    t = doc.add_table(rows=5, cols=2)
    add_header_row(t, ['Table', 'Primary Key Columns'])
    for i, row_data in enumerate(keys):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 6. EXTRACTION LOGIC
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('6. Extraction Logic', level=1)

    doc.add_heading('6.1 Date Range Resolution', level=2)
    doc.add_paragraph(
        'The pipeline resolves the extraction date range using the following priority:\n\n'
        '1. Explicit start_date / end_date arguments (highest priority)\n'
        '2. Full refresh mode: extract last max_months x 30 days (default 12 months = 360 days)\n'
        '3. Incremental mode (default): extract last lookback_days (default 30 from config)\n\n'
        'Dates are converted to Unix timestamps for the Threads API. The since parameter is '
        'inclusive (start of start_date) and until is exclusive (start of end_date + 1 day).'
    )

    doc.add_heading('6.2 Profiles Extraction', level=2)
    doc.add_paragraph(
        'Calls GET /me for each account. Returns the authenticated user\'s profile. '
        'No date range needed. Extracted on every run to keep profile data current.'
    )

    doc.add_heading('6.3 Posts Extraction', level=2)
    doc.add_paragraph(
        'Calls GET /{user_id}/threads with since/until date filters and cursor pagination. '
        'Retrieves all posts within the date range, up to 100 per page. In test mode, '
        'extraction stops after 25 records per user. Timestamps are normalized via '
        'parse_facebook_timestamp().'
    )

    doc.add_heading('6.4 Post Insights Extraction', level=2)
    doc.add_paragraph(
        'For each post extracted, calls GET /{thread_id}/insights to retrieve engagement metrics. '
        'Posts with media_type = REPOST_FACADE are skipped because the API does not support '
        'insights for reposts. The response is flattened from the nested metric/values format '
        'into a single row per post.'
    )

    doc.add_heading('6.5 User Insights Extraction', level=2)
    doc.add_paragraph(
        'Calls GET /{user_id}/threads_insights with period=day. The date range is chunked into '
        '90-day windows to avoid API timeouts. For each chunk, the response is pivoted from '
        'metric-per-day format into one row per day with all metrics as columns. Usernames are '
        'enriched from a profiles cache built during the profiles extraction step.'
    )

    doc.add_heading('6.6 Record Processing', level=2)
    doc.add_paragraph(
        'All records pass through _process_records_to_parquet() which:\n\n'
        '  1. Adds execution metadata (execution_id, source, extracted_at)\n'
        '  2. Normalizes timestamp fields to ISO 8601 format\n'
        '  3. Converts all ID values to strings for BigQuery STRING schema compatibility\n'
        '  4. Serializes nested dicts/lists to JSON strings\n'
        '  5. Writes to a local Parquet file via shared/gcs_pipeline.extract_to_local_parquet()'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 7. DATA TRANSFORMATIONS (SQL)
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('7. Data Transformations (SQL)', level=1)

    doc.add_heading('7.1 threads_post_insights_daily.sql', level=2)
    doc.add_paragraph(
        'Post-process SQL that creates BN_Warehouse.threads_post_insights_daily. '
        'Executed automatically after each pipeline run.'
    )

    doc.add_heading('Join Logic', level=3)
    doc.add_paragraph(
        'LEFT JOIN between threads_data.threads_posts and threads_data.threads_post_insights '
        'on thread_id and threads_user_id. This preserves posts that may not yet have insights '
        '(newly published content).'
    )

    doc.add_heading('Metric Casting', level=3)
    doc.add_paragraph(
        'All metrics are stored as STRING in staging/production tables and cast to FLOAT64 in '
        'the analytics layer. IFNULL wraps each SAFE_CAST to default NULL metrics to 0.0.'
    )

    doc.add_heading('Derived Metrics', level=3)
    items = [
        'total_engagements = likes + replies + reposts + quotes + shares',
        'engagement_rate = total_engagements / views (using SAFE_DIVIDE to avoid division by zero)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Table Configuration', level=3)
    items = [
        'Partitioned by date (DATE type, derived from post timestamp)',
        'Clustered by threads_user_id, thread_id',
        'Requires partition filter for query performance',
        'Labels: source=threads, pipeline=bi, grain=daily',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 8. RATE LIMITING & RETRY STRATEGY
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('8. Rate Limiting & Retry Strategy', level=1)

    doc.add_heading('8.1 Rate Limits', level=2)
    limits = [
        ('default', '200 calls/hour', '18 seconds', 'Profile, post listing'),
        ('insights', '100 calls/hour', '36 seconds', 'Post insights, user insights'),
        ('posts', '150 calls/hour', '24 seconds', 'Post content retrieval'),
    ]
    t = doc.add_table(rows=4, cols=4)
    add_header_row(t, ['Endpoint Type', 'Limit', 'Min Interval', 'Used By'])
    for i, row_data in enumerate(limits):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_paragraph(
        '\nA hardcoded minimum delay of 1.5 seconds is enforced regardless of the calculated '
        'interval. The wait_for_threads() function is called before every API request.'
    )

    doc.add_heading('8.2 Retry Strategy', level=2)
    doc.add_paragraph('The API client implements a three-tier retry strategy:')

    retries = [
        ('HTTP 429 (Rate Limited)', 'Wait for Retry-After header duration, then retry', '3'),
        ('HTTP 5xx (Server Error)', 'Exponential backoff: 4s, 8s, 16s', '3'),
        ('Connection/Timeout Errors', 'Exponential backoff: 4s, 8s, 16s', '3'),
    ]
    t = doc.add_table(rows=4, cols=3)
    add_header_row(t, ['Error Type', 'Behavior', 'Max Retries'])
    for i, row_data in enumerate(retries):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('8.3 Session Configuration', level=2)
    doc.add_paragraph(
        'The requests.Session is configured with an HTTPAdapter using urllib3.Retry:\n\n'
        '  - Total retries: 3\n'
        '  - Backoff factor: 1 (exponential: 1s, 2s, 4s)\n'
        '  - Retry on status codes: 429, 500, 502, 503, 504\n'
        '  - Retry on methods: GET, POST\n'
        '  - Connection timeout: 30 seconds'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 9. MULTI-ACCOUNT ARCHITECTURE
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('9. Multi-Account Architecture', level=1)

    doc.add_heading('9.1 Design Rationale', level=2)
    doc.add_paragraph(
        'Unlike Facebook/Instagram where a single Page token can access multiple assets, '
        'the Threads API is strictly per-user. Each Threads account requires its own OAuth '
        'token. Bionews operates 12+ Threads accounts, necessitating a multi-account token '
        'management system.'
    )

    doc.add_heading('9.2 Account Discovery', level=2)
    doc.add_paragraph(
        'Threads accounts are linked to Instagram Business accounts. The list_threads_accounts.py '
        'utility discovers all linked Threads accounts by:\n\n'
        '  1. Querying Facebook API for all managed Pages\n'
        '  2. For each Page, checking for an instagram_business_account\n'
        '  3. Fetching the Instagram username (which matches the Threads username)\n'
        '  4. Outputting URLs: https://www.threads.net/@{username}'
    )

    doc.add_heading('9.3 Extraction Flow', level=2)
    steps = [
        'Load all accounts from .threads_tokens.json',
        'For each account: auto-refresh token if within 7 days of expiry',
        'For each account: verify token validity via GET /me',
        'For each account + table: extract data and accumulate records',
        'Posts are cached in memory for subsequent post_insights extraction',
        'Write one Parquet file per table (aggregated across all accounts)',
        'Each row is tagged with threads_user_id for tenant isolation',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {s}', style='List Number')

    doc.add_heading('9.4 Tenant Isolation', level=2)
    doc.add_paragraph(
        'Every table includes a threads_user_id column that identifies the source account. '
        'This serves as the tenant_field in the configuration, enabling per-account queries '
        'and multi-tenant analytics without cross-contamination.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 10. CONFIGURATION REFERENCE
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('10. Configuration Reference', level=1)

    doc.add_paragraph('Configuration file: configs/threads.yaml')

    doc.add_heading('10.1 Source Connection', level=2)
    cfg = [
        ('source.name', 'threads', 'Pipeline name'),
        ('source.type', 'threads', 'Plugin type (registered in config_loader.py)'),
        ('source.connection.access_token', '${THREADS_ACCESS_TOKEN}', 'Legacy single-account token'),
        ('source.connection.app_id', '${THREADS_APP_ID}', 'OAuth app ID'),
        ('source.connection.app_secret', '${THREADS_APP_SECRET}', 'OAuth app secret'),
    ]
    t = doc.add_table(rows=len(cfg) + 1, cols=3)
    add_header_row(t, ['Key', 'Default', 'Description'])
    for i, row_data in enumerate(cfg):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('10.2 Pipeline Settings', level=2)
    cfg = [
        ('pipeline.staging_dataset', 'threads_staging', 'BigQuery staging dataset'),
        ('pipeline.production_dataset', 'threads_data', 'BigQuery production dataset'),
        ('pipeline.test_dataset', 'threads_test', 'BigQuery test dataset'),
        ('pipeline.gcs_bucket', '${GCS_BUCKET}', 'GCS bucket for Parquet files'),
        ('pipeline.gcs_path', 'threads/raw', 'GCS path for raw extracts'),
        ('pipeline.parallel_workers', '2', 'Concurrent extraction threads'),
        ('pipeline.batch_size', '500', 'Records per batch'),
    ]
    t = doc.add_table(rows=len(cfg) + 1, cols=3)
    add_header_row(t, ['Key', 'Default', 'Description'])
    for i, row_data in enumerate(cfg):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('10.3 Rate Limits', level=2)
    cfg = [
        ('source.rate_limits.default', '200', 'General API calls per hour'),
        ('source.rate_limits.insights', '100', 'Insights API calls per hour'),
        ('source.rate_limits.posts', '150', 'Posts API calls per hour'),
    ]
    t = doc.add_table(rows=len(cfg) + 1, cols=3)
    add_header_row(t, ['Key', 'Default', 'Description'])
    for i, row_data in enumerate(cfg):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('10.4 Extractor Settings', level=2)
    cfg = [
        ('extractor.lookback_days', '30', 'Default incremental lookback period'),
        ('extractor.full_refresh.max_months', '12', 'Full refresh history depth'),
        ('extractor.full_refresh.auto_truncate', 'true', 'Truncate before full refresh'),
    ]
    t = doc.add_table(rows=len(cfg) + 1, cols=3)
    add_header_row(t, ['Key', 'Default', 'Description'])
    for i, row_data in enumerate(cfg):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 11. SCHEDULING & OPERATIONS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('11. Scheduling & Operations', level=1)

    doc.add_heading('11.1 Pipeline Schedule', level=2)
    cfg = [
        ('Schedule', '0 4 * * * (daily at 4:00 AM)'),
        ('Timezone', 'America/New_York'),
        ('Retry on failure', 'Yes'),
        ('Max retries', '3'),
        ('Retry delay', '30 minutes'),
        ('Timeout', '60 minutes'),
    ]
    t = doc.add_table(rows=len(cfg) + 1, cols=2)
    add_header_row(t, ['Setting', 'Value'])
    for i, row_data in enumerate(cfg):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_heading('11.2 Token Refresh Schedule', level=2)
    doc.add_paragraph(
        'Tokens should be refreshed every 50 days (well before the 60-day expiry). '
        'The pipeline auto-refreshes tokens within 7 days of expiry, but a scheduled '
        'refresh provides an additional safety margin.'
    )
    doc.add_paragraph('Linux cron:')
    doc.add_paragraph(
        '  0 6 */50 * * cd /path && python shared/threads_token_manager.py --refresh'
    )
    doc.add_paragraph('Windows Task Scheduler:')
    doc.add_paragraph(
        '  schtasks /create /tn "Threads Token Refresh" /tr "python '
        'c:\\\\orchestrator\\\\shared\\\\threads_token_manager.py --refresh" /sc daily /mo 50'
    )

    doc.add_heading('11.3 Running the Pipeline', level=2)
    runs = [
        ('Incremental (default)', 'python orchestrate.py --source threads'),
        ('Full refresh', 'python orchestrate.py --source threads --refresh-mode full'),
        ('Specific tables', 'python orchestrate.py --source threads --tables posts,post_insights'),
        ('Date range', 'python orchestrate.py --source threads --start-date 2026-01-01 --end-date 2026-01-31'),
        ('Test mode', 'python orchestrate.py --source threads --test'),
    ]
    t = doc.add_table(rows=len(runs) + 1, cols=2)
    add_header_row(t, ['Mode', 'Command'])
    for i, row_data in enumerate(runs):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 12. ERROR HANDLING
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('12. Error Handling', level=1)

    doc.add_heading('12.1 Error Categories', level=2)
    errors = [
        ('Token expired', 'Auto-refresh if within 7 days; skip account if fully expired', 'WARNING'),
        ('Token invalid', 'GET /me fails; skip account, continue with others', 'ERROR'),
        ('Rate limited (429)', 'Wait for Retry-After header, retry up to 3 times', 'WARNING'),
        ('Server error (5xx)', 'Exponential backoff, retry up to 3 times', 'WARNING'),
        ('REPOST_FACADE insights', 'Silently skip — this media type does not support insights', 'DEBUG'),
        ('Missing insights', 'Log at DEBUG level, skip record, continue extraction', 'DEBUG'),
        ('No accounts configured', 'Raise exception — pipeline cannot proceed', 'CRITICAL'),
        ('Chunk extraction failure', 'Log warning, continue with remaining chunks', 'WARNING'),
    ]
    t = doc.add_table(rows=len(errors) + 1, cols=3)
    add_header_row(t, ['Scenario', 'Behavior', 'Log Level'])
    for i, row_data in enumerate(errors):
        add_data_row(t, i + 1, row_data)
    set_table_style(t)

    doc.add_heading('12.2 Graceful Degradation', level=2)
    doc.add_paragraph(
        'The pipeline is designed for maximum data collection even when individual components fail:\n\n'
        '  - Per-account isolation: if one account\'s token fails, others continue\n'
        '  - Per-table isolation: if one table extraction fails, others complete\n'
        '  - Per-chunk isolation: if one date chunk fails, remaining chunks are processed\n'
        '  - Per-post isolation: if insights for one post fail, other posts are unaffected\n\n'
        'The pipeline returns a summary with total_rows, successful tables, and failed_tables counts.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 13. TESTING & APP REVIEW
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('13. Testing & App Review', level=1)

    doc.add_heading('13.1 Integration Test Suite', level=2)
    doc.add_paragraph(
        'File: tests/integration/test_threads_app_review.py (615 lines)\n\n'
        'Purpose: Demonstrates all three Threads permissions for Meta App Review submission. '
        'The test suite exercises each permission scope independently and produces a '
        'color-coded PASS/FAIL/SKIP report.'
    )

    doc.add_heading('13.2 Permission Tests', level=2)
    tests = [
        ('threads_basic', 'GET /me (profile), GET /me/threads (posts)', 'Profile fields, post listing with metadata'),
        ('threads_manage_insights', 'GET /{id}/insights, GET /me/threads_insights', 'Post metrics, daily account metrics'),
        ('threads_read_replies', 'GET /{id}/replies, GET /{id}/conversation', 'Direct replies, nested conversations'),
    ]
    t = doc.add_table(rows=4, cols=3)
    add_header_row(t, ['Permission', 'API Calls', 'Validates'])
    for i, row_data in enumerate(tests):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_heading('13.3 Running Tests', level=2)
    doc.add_paragraph(
        '  # All permissions\n'
        '  python tests/integration/test_threads_app_review.py\n\n'
        '  # Single permission\n'
        '  python tests/integration/test_threads_app_review.py --permission threads_basic\n\n'
        '  # Custom token\n'
        '  python tests/integration/test_threads_app_review.py --token THAA...'
    )

    doc.add_heading('13.4 Meta App Setup', level=2)
    doc.add_paragraph(
        'The Threads API requires a Meta app created via the new "use-case" flow:\n\n'
        '  1. Navigate to https://developers.facebook.com/apps/creation/\n'
        '  2. Select "Access the Threads API" use case\n'
        '  3. Do NOT connect a business portfolio initially\n'
        '  4. Add permissions: threads_basic, threads_manage_insights, threads_read_replies\n'
        '  5. Configure OAuth redirect: https://localhost:8899/callback (HTTPS required)\n'
        '  6. In Development mode: add up to 50 Threads Testers\n\n'
        'Important: The old "product-based" app creation flow does NOT offer Threads API access. '
        'Apps created with the old flow cannot be upgraded.'
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 14. KNOWN LIMITATIONS & GOTCHAS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('14. Known Limitations & Gotchas', level=1)

    gotchas = [
        ('No SDK support', 'The facebook_business Python SDK does not support the Threads API. All calls use raw requests.'),
        ('Per-user tokens', 'Each Threads account requires its own OAuth token. There is no Pages or Business Manager concept for Threads.'),
        ('REPOST_FACADE', 'The REPOST_FACADE media type does not support insights. The pipeline silently skips these posts during insights extraction.'),
        ('followers_count threshold', 'The followers_count and follower_demographics metrics require the account to have 100 or more followers. Accounts below this threshold will not return follower data.'),
        ('60-day token expiry', 'Long-lived tokens expire after 60 days. The pipeline auto-refreshes within 7 days, but a scheduled refresh every 50 days is recommended as a safety net.'),
        ('localhost redirect rejected', 'Meta rejects localhost redirect URLs for Threads OAuth. Use the User Token Generator (developers.facebook.com) instead, then add tokens via --add-token.'),
        ('appsecret_proof errors', 'The facebook_business SDK auto-attaches appsecret_proof to requests, which causes errors on the Threads API. The pipeline uses raw requests to avoid this.'),
        ('Old app flow incompatible', 'Meta apps created with the old "product-based" flow cannot access the Threads API. A new app must be created via the "use-case" flow.'),
        ('Development mode limits', 'In development mode (before App Review), only Threads Testers (up to 50) can authorize. Production use requires passing App Review.'),
        ('90-day chunk limit', 'User insights are chunked into 90-day windows. Larger date ranges are automatically split to avoid API timeouts.'),
    ]
    for title, desc in gotchas:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════
    # 15. APPENDIX: ENVIRONMENT VARIABLES
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading('15. Appendix: Environment Variables', level=1)

    envvars = [
        ('THREADS_ACCESS_TOKEN', 'Yes (if no token store)', 'Long-lived OAuth token (legacy single-account)'),
        ('THREADS_USER_ID', 'Yes (if no token store)', 'Threads user ID (legacy single-account)'),
        ('THREADS_APP_ID', 'Yes', 'Meta app client ID (from Threads section of app dashboard)'),
        ('THREADS_APP_SECRET', 'Yes', 'Meta app client secret'),
        ('THREADS_REDIRECT_URI', 'No', 'Custom OAuth redirect URI (default: https://localhost:8899/callback)'),
        ('GCP_PROJECT_ID', 'No', 'BigQuery project ID (default: bi-data-391216)'),
        ('GCS_BUCKET', 'Yes', 'Google Cloud Storage bucket for Parquet files'),
        ('ALERT_RECIPIENT_EMAIL', 'No', 'Email for pipeline failure notifications'),
        ('SMTP_HOST', 'No', 'SMTP server for email alerts'),
        ('SMTP_USER', 'No', 'SMTP username'),
    ]
    t = doc.add_table(rows=len(envvars) + 1, cols=3)
    add_header_row(t, ['Variable', 'Required', 'Description'])
    for i, row_data in enumerate(envvars):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    doc.add_heading('15.1 Files', level=2)
    files = [
        ('.threads_tokens.json', 'Multi-account token store (created by token manager)'),
        ('.threads_token_meta.json', 'Legacy single-account token metadata'),
        ('.env', 'Environment variables (fallback for single-account)'),
        ('configs/threads.yaml', 'Pipeline configuration'),
    ]
    t = doc.add_table(rows=len(files) + 1, cols=2)
    add_header_row(t, ['File', 'Purpose'])
    for i, row_data in enumerate(files):
        add_data_row(t, i + 1, row_data, bold_first=True)
    set_table_style(t)

    # ── Save ──
    out_path = os.path.join(os.path.dirname(__file__), 'Threads_Pipeline_Specification.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')
    return out_path

if __name__ == '__main__':
    build_doc()
