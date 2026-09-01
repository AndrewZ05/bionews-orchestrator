"""Generate Instagram Data Dictionary Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# -- Styles --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# -- Helper functions --
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return table

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
title = doc.add_heading('Instagram Analytics Pipeline', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
subtitle = doc.add_heading('Data Dictionary & Field Reference', level=1)
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0x3D, 0x6B, 0x8E)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Prepared for: ').bold = True
p.add_run('BioNews Senior Leadership')
p = doc.add_paragraph()
p.add_run('Date: ').bold = True
p.add_run('February 2026')
p = doc.add_paragraph()
p.add_run('Source: ').bold = True
p.add_run('Meta Instagram Graph API v25.0')
p = doc.add_paragraph()
p.add_run('Accounts: ').bold = True
p.add_run('35 BioNews Instagram Business accounts')

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Data Architecture',
    '3. Source Tables (instagram_data)',
    '   3.1 instagram_accounts',
    '   3.2 instagram_media',
    '   3.3 instagram_media_insights',
    '   3.4 instagram_account_insights',
    '   3.5 instagram_stories',
    '   3.6 instagram_story_insights',
    '4. Report Tables (BN_Warehouse)',
    '   4.1 instagram_media_insights_daily',
    '   4.2 instagram_account_insights_daily',
    '5. Data Refresh Schedule',
    '6. Known Limitations & API Constraints',
    '7. Upcoming Changes',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. OVERVIEW
# ============================================================
add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The Instagram Analytics Pipeline extracts data from 35 BioNews Instagram Business accounts '
    'via Meta\'s Instagram Graph API (v25.0) and loads it into Google BigQuery for reporting. '
    'Data is refreshed daily via automated orchestration on a Linux server.'
)
doc.add_paragraph(
    'The pipeline produces 6 source tables in the instagram_data dataset (raw API data) and '
    '2 report tables in the BN_Warehouse dataset (enriched, analysis-ready data). This document '
    'provides a complete field-level reference for every column in every table.'
)

# ============================================================
# 2. DATA ARCHITECTURE
# ============================================================
add_heading('2. Data Architecture', level=1)
doc.add_paragraph(
    'Data flows through the following stages:'
)
arch_items = [
    ('Extraction', 'instagram_extractor.py calls the Instagram Graph API for each of the 35 accounts.'),
    ('Staging', 'Raw data is written to Parquet files, uploaded to Google Cloud Storage, then loaded into BigQuery staging tables.'),
    ('Production', 'A hash-merge process compares staging rows to production rows and upserts only changed/new records into instagram_data.* tables.'),
    ('Reporting', 'Post-process SQL scripts join and aggregate source tables into BN_Warehouse.*_daily report tables, partitioned by date.'),
]
for label, desc in arch_items:
    p = doc.add_paragraph()
    p.add_run(f'{label}: ').bold = True
    p.add_run(desc)

add_heading('Table Summary', level=2)
add_table(
    ['Dataset', 'Table', 'Description', 'Grain', 'Row Count'],
    [
        ['instagram_data', 'instagram_accounts', 'Profile metadata for all 35 IG Business accounts', 'One row per account', '35'],
        ['instagram_data', 'instagram_media', 'Every post (images, videos, carousels, reels) with native counts', 'One row per post', '~24,000'],
        ['instagram_data', 'instagram_media_insights', 'Per-post performance metrics from the Insights API', 'One row per post', '~16,000'],
        ['instagram_data', 'instagram_account_insights', 'Account-level daily performance metrics (reach, followers)', 'One row per account per day', '~26,000'],
        ['instagram_data', 'instagram_stories', 'Currently active stories (ephemeral, 24-hour lifespan)', 'One row per story', '~160'],
        ['instagram_data', 'instagram_story_insights', 'Per-story engagement metrics', 'One row per story', '~130'],
        ['BN_Warehouse', 'instagram_media_insights_daily', 'Joined media + insights with derived engagement rate', 'One row per post', 'Derived'],
        ['BN_Warehouse', 'instagram_account_insights_daily', 'Daily account performance with media-aggregated engagement', 'One row per account per day', 'Derived'],
    ],
    col_widths=[3.5, 5.5, 7, 4, 2.5]
)

doc.add_page_break()

# ============================================================
# 3. SOURCE TABLES
# ============================================================
add_heading('3. Source Tables (instagram_data)', level=1)
doc.add_paragraph(
    'Source tables contain raw data extracted directly from the Instagram Graph API. '
    'These tables are loaded incrementally and serve as the foundation for all downstream reporting.'
)

# 3.1 instagram_accounts
add_heading('3.1 instagram_accounts', level=2)
doc.add_paragraph(
    'Profile metadata for all 35 Instagram Business accounts linked to BioNews Facebook Pages. '
    'Fully refreshed daily. This table provides the master list of accounts and their current profile information.'
)
p = doc.add_paragraph()
p.add_run('Refresh: ').bold = True
p.add_run('Daily full refresh')
p = doc.add_paragraph()
p.add_run('Grain: ').bold = True
p.add_run('One row per Instagram Business account')

add_table(
    ['Column', 'Data Type', 'Description', 'Example'],
    [
        ['ig_user_id', 'STRING', 'Unique Instagram Business account identifier assigned by Meta. Primary key for joining to other tables.', '17841400123456789'],
        ['username', 'STRING', 'Instagram handle (without @). The public username visible on Instagram.', 'msnewstoday'],
        ['name', 'STRING', 'Display name shown on the Instagram profile. Can differ from username.', 'MS News Today'],
        ['biography', 'STRING', 'Profile bio text. May contain line breaks, emojis, and hashtags. Limited to 150 characters by Instagram.', 'Your source for MS news and resources.'],
        ['website', 'STRING', 'Website URL displayed on the profile. May be NULL if not set.', 'https://www.msnewstoday.com'],
        ['profile_picture_url', 'STRING', 'URL to the profile image. URL may expire; primarily for reference.', 'https://scontent.cdninstagram.com/...'],
        ['followers_count', 'STRING', 'Total number of followers at time of extraction. Stored as STRING for schema consistency.', '15234'],
        ['follows_count', 'STRING', 'Total number of accounts this account follows.', '127'],
        ['media_count', 'STRING', 'Total number of posts (all types) published by the account.', '1842'],
        ['page_id', 'STRING', 'Facebook Page ID linked to this Instagram Business account. Required for API access.', '123456789012345'],
        ['page_name', 'STRING', 'Name of the linked Facebook Page.', 'MS News Today'],
    ],
    col_widths=[3.5, 2.5, 8, 4]
)

# 3.2 instagram_media
add_heading('3.2 instagram_media', level=2)
doc.add_paragraph(
    'Every post published by the 35 accounts, including images, videos, carousels, and reels. '
    'Extracted with a 30-day incremental lookback to capture edits and updated native counts. '
    'Historical data goes back to 2016 for long-running accounts.'
)
p = doc.add_paragraph()
p.add_run('Refresh: ').bold = True
p.add_run('Daily incremental (30-day lookback window)')
p = doc.add_paragraph()
p.add_run('Grain: ').bold = True
p.add_run('One row per post')

add_table(
    ['Column', 'Data Type', 'Description', 'Example'],
    [
        ['id', 'STRING', 'Unique post identifier assigned by Instagram. Primary key. Used to join to instagram_media_insights.', '17854360229XXXXXX'],
        ['ig_user_id', 'STRING', 'Instagram Business account ID that published this post. Foreign key to instagram_accounts.', '17841400123456789'],
        ['caption', 'STRING', 'Full caption text of the post. May contain emojis, hashtags, and mentions. Can be NULL for posts without captions.', 'New research shows...'],
        ['media_type', 'STRING', 'Type of media. IMAGE = single photo, VIDEO = single video or reel, CAROUSEL_ALBUM = multi-image/video swipe post.', 'IMAGE'],
        ['media_product_type', 'STRING', 'Product surface where the media was published. FEED = standard post, REELS = short-form video, STORY = ephemeral content.', 'FEED'],
        ['media_url', 'STRING', 'URL to the media asset. May expire after time. Not reliable for long-term storage.', 'https://scontent.cdninstagram.com/...'],
        ['permalink', 'STRING', 'Permanent public link to the post on Instagram. Does not expire.', 'https://www.instagram.com/p/ABC123/'],
        ['shortcode', 'STRING', 'Short alphanumeric code used in Instagram URLs. Extracted from the permalink.', 'ABC123'],
        ['timestamp', 'STRING', 'ISO 8601 date/time when the post was published. Stored as STRING; cast to TIMESTAMP in report SQL.', '2026-02-15T14:30:00+0000'],
        ['like_count', 'STRING', 'Native like count from the post object. Available for all media types including carousels. Updated on each extraction.', '245'],
        ['comments_count', 'STRING', 'Native comment count from the post object. Available for all media types. Updated on each extraction.', '18'],
    ],
    col_widths=[3.5, 2.5, 8, 4]
)

# 3.3 instagram_media_insights
add_heading('3.3 instagram_media_insights', level=2)
doc.add_paragraph(
    'Performance metrics for each post, sourced from Instagram\'s Insights API. Provides deeper analytics '
    'than native counts (e.g., reach, shares, saves). These metrics are lifetime values that accumulate over time.'
)
p = doc.add_paragraph()
p.add_run('Refresh: ').bold = True
p.add_run('Daily incremental (30-day lookback window)')
p = doc.add_paragraph()
p.add_run('Grain: ').bold = True
p.add_run('One row per post')
p = doc.add_paragraph()
p.add_run('Important: ').bold = True
p.add_run('CAROUSEL_ALBUM posts return 0 for likes, comments, and shares through the Insights API. '
           'Use native like_count and comments_count from instagram_media for carousel engagement.')

add_table(
    ['Column', 'Data Type', 'Description', 'Example'],
    [
        ['media_id', 'STRING', 'Post identifier. Foreign key joining to instagram_media.id.', '17854360229XXXXXX'],
        ['ig_user_id', 'STRING', 'Instagram Business account ID. Foreign key to instagram_accounts.', '17841400123456789'],
        ['media_type', 'STRING', 'Type of media (IMAGE, VIDEO, CAROUSEL_ALBUM). Denormalized from media table for convenience.', 'VIDEO'],
        ['views', 'STRING', 'Total number of times the post was displayed on screen, including repeat views by the same user. Replaces deprecated "impressions" metric (April 2025).', '3847'],
        ['reach', 'STRING', 'Number of unique accounts that saw the post at least once. Always less than or equal to views.', '2156'],
        ['saved', 'STRING', 'Number of times the post was saved/bookmarked by users.', '42'],
        ['likes', 'STRING', 'Like count from Insights API. Returns 0 for CAROUSEL_ALBUM posts (API limitation).', '185'],
        ['comments', 'STRING', 'Comment count from Insights API. Returns 0 for CAROUSEL_ALBUM posts (API limitation).', '12'],
        ['shares', 'STRING', 'Number of times the post was shared (DM, stories, external). Returns 0 for CAROUSEL_ALBUM posts.', '28'],
        ['total_interactions', 'STRING', 'Sum of all engagement actions (likes + comments + shares + saves + other interactions). Provided directly by the API.', '267'],
    ],
    col_widths=[3.5, 2.5, 8, 4]
)

# 3.4 instagram_account_insights
add_heading('3.4 instagram_account_insights', level=2)
doc.add_paragraph(
    'Account-level daily performance metrics extracted from the Instagram Insights API. '
    'Provides reach (unique accounts that saw content) and follower growth data. '
    'Views and engagement metrics are NOT available at the account level with daily granularity '
    'due to API limitations; these are derived from media-level data in the report tables.'
)
p = doc.add_paragraph()
p.add_run('Refresh: ').bold = True
p.add_run('Daily incremental (30-day chunks for reach; follower_count limited to last 30 days by API)')
p = doc.add_paragraph()
p.add_run('Grain: ').bold = True
p.add_run('One row per account per day')

add_table(
    ['Column', 'Data Type', 'Description', 'Example'],
    [
        ['ig_user_id', 'STRING', 'Instagram Business account identifier. Primary key (with date).', '17841400123456789'],
        ['username', 'STRING', 'Instagram handle at time of extraction.', 'msnewstoday'],
        ['account_id', 'STRING', 'Alternative account identifier used in some API contexts.', '17841400123456789'],
        ['page_id', 'STRING', 'Linked Facebook Page ID.', '123456789012345'],
        ['date', 'STRING', 'Calendar date of the metric in YYYY-MM-DD format.', '2026-02-15'],
        ['reach', 'STRING', 'Number of unique accounts that saw ANY content from this account on this date. Sourced from API period=day.', '4523'],
        ['follower_count', 'STRING', 'Net new followers gained on this date (NOT cumulative total). Only available for the last 30 days; older dates show 0. Once captured, data is preserved.', '12'],
        ['views', 'STRING', 'Always "0" in raw table. Account-level views only supports aggregate responses from the API, not daily breakdowns. Daily views are derived from media data in the report table.', '0'],
    ],
    col_widths=[3.5, 2.5, 8, 4]
)

# 3.5 instagram_stories
add_heading('3.5 instagram_stories', level=2)
doc.add_paragraph(
    'Currently active Instagram Stories. Stories are ephemeral content that expires after 24 hours. '
    'The API only returns stories that are currently live, so this table captures whatever is active at '
    'the time of each daily pipeline run. Historical stories cannot be recovered once expired.'
)
p = doc.add_paragraph()
p.add_run('Refresh: ').bold = True
p.add_run('Daily full refresh (captures 24-hour window)')
p = doc.add_paragraph()
p.add_run('Grain: ').bold = True
p.add_run('One row per active story')

add_table(
    ['Column', 'Data Type', 'Description', 'Example'],
    [
        ['id', 'STRING', 'Unique story identifier assigned by Instagram. Primary key. Used to join to instagram_story_insights.', '17854360229XXXXXX'],
        ['ig_user_id', 'STRING', 'Instagram Business account ID that published the story. Foreign key to instagram_accounts.', '17841400123456789'],
        ['caption', 'STRING', 'Story caption text. Often NULL as most stories use visual elements rather than text captions.', 'NULL'],
        ['media_type', 'STRING', 'Type of story media. IMAGE = photo story, VIDEO = video story.', 'IMAGE'],
        ['media_url', 'STRING', 'URL to the story media asset. Expires after the story is removed.', 'https://scontent.cdninstagram.com/...'],
        ['permalink', 'STRING', 'Permanent link to the story on Instagram. May become invalid after story expires.', 'https://www.instagram.com/stories/...'],
        ['timestamp', 'STRING', 'ISO 8601 date/time when the story was published.', '2026-02-15T09:00:00+0000'],
    ],
    col_widths=[3.5, 2.5, 8, 4]
)

# 3.6 instagram_story_insights
add_heading('3.6 instagram_story_insights', level=2)
doc.add_paragraph(
    'Engagement metrics for active stories. Only available while the story is live (24-hour window). '
    'Several story metrics were deprecated by Meta in January 2025 (navigation, exits, saved, total_interactions), '
    'leaving reach and replies as the available metrics.'
)
p = doc.add_paragraph()
p.add_run('Refresh: ').bold = True
p.add_run('Daily full refresh')
p = doc.add_paragraph()
p.add_run('Grain: ').bold = True
p.add_run('One row per active story')

add_table(
    ['Column', 'Data Type', 'Description', 'Example'],
    [
        ['story_id', 'STRING', 'Story identifier. Foreign key joining to instagram_stories.id.', '17854360229XXXXXX'],
        ['ig_user_id', 'STRING', 'Instagram Business account ID. Foreign key to instagram_accounts.', '17841400123456789'],
        ['reach', 'STRING', 'Number of unique accounts that viewed this story.', '1234'],
        ['replies', 'STRING', 'Number of direct message replies sent in response to this story.', '3'],
    ],
    col_widths=[3.5, 2.5, 8, 4]
)

doc.add_page_break()

# ============================================================
# 4. REPORT TABLES
# ============================================================
add_heading('4. Report Tables (BN_Warehouse)', level=1)
doc.add_paragraph(
    'Report tables are rebuilt daily from source tables using post-process SQL. They join multiple '
    'source tables, cast data types, compute derived metrics, and provide analysis-ready data. '
    'All report tables are partitioned by date and clustered by ig_user_id for efficient querying.'
)

# 4.1 instagram_media_insights_daily
add_heading('4.1 instagram_media_insights_daily', level=2)
doc.add_paragraph(
    'Joined view of media posts and their Insights API metrics, with derived engagement rate. '
    'Combines data from instagram_media and instagram_media_insights into a single analysis-ready table.'
)
p = doc.add_paragraph()
p.add_run('Source tables: ').bold = True
p.add_run('instagram_media + instagram_media_insights')
p = doc.add_paragraph()
p.add_run('Partitioned by: ').bold = True
p.add_run('date (derived from post timestamp)')
p = doc.add_paragraph()
p.add_run('Clustered by: ').bold = True
p.add_run('ig_user_id')

add_table(
    ['Column', 'Data Type', 'Source', 'Description'],
    [
        ['id', 'STRING', 'instagram_media', 'Unique post identifier'],
        ['ig_user_id', 'STRING', 'instagram_media', 'Instagram Business account ID'],
        ['username', 'STRING', 'instagram_accounts', 'Instagram handle'],
        ['caption', 'STRING', 'instagram_media', 'Post caption text'],
        ['media_type', 'STRING', 'instagram_media', 'IMAGE, VIDEO, or CAROUSEL_ALBUM'],
        ['media_product_type', 'STRING', 'instagram_media', 'FEED, REELS, or STORY'],
        ['permalink', 'STRING', 'instagram_media', 'Permanent link to the post'],
        ['date', 'DATE', 'Derived', 'Date post was published (cast from timestamp)'],
        ['like_count', 'FLOAT64', 'instagram_media', 'Native like count from post object'],
        ['comments_count', 'FLOAT64', 'instagram_media', 'Native comment count from post object'],
        ['views', 'FLOAT64', 'instagram_media_insights', 'Total times post was displayed (lifetime)'],
        ['reach', 'FLOAT64', 'instagram_media_insights', 'Unique accounts that saw the post'],
        ['saved', 'FLOAT64', 'instagram_media_insights', 'Times post was bookmarked'],
        ['likes', 'FLOAT64', 'instagram_media_insights', 'Likes from Insights API (0 for carousels)'],
        ['comments', 'FLOAT64', 'instagram_media_insights', 'Comments from Insights API (0 for carousels)'],
        ['shares', 'FLOAT64', 'instagram_media_insights', 'Times post was shared (0 for carousels)'],
        ['total_interactions', 'FLOAT64', 'instagram_media_insights', 'Sum of all engagement actions'],
        ['engagement_rate', 'FLOAT64', 'Derived', 'total_interactions / reach (NULL if reach = 0)'],
    ],
    col_widths=[3.5, 2, 4, 7]
)

# 4.2 instagram_account_insights_daily
add_heading('4.2 instagram_account_insights_daily', level=2)
doc.add_paragraph(
    'Daily account-level performance combining API metrics (reach, follower growth) with '
    'engagement metrics aggregated from individual post data. This is the primary table for '
    'account-level trend analysis and cross-account comparison.'
)
p = doc.add_paragraph()
p.add_run('Source tables: ').bold = True
p.add_run('instagram_account_insights + instagram_media + instagram_media_insights')
p = doc.add_paragraph()
p.add_run('Partitioned by: ').bold = True
p.add_run('date')
p = doc.add_paragraph()
p.add_run('Clustered by: ').bold = True
p.add_run('ig_user_id')
p = doc.add_paragraph()
p.add_run('Join logic: ').bold = True
p.add_run('Account insights (reach, followers from API) LEFT JOIN media insights aggregated by '
           'account + publish date. Days with no posts will show views/engagement = 0 and posts_published = 0.')

add_table(
    ['Column', 'Data Type', 'Source', 'Description'],
    [
        ['ig_user_id', 'STRING', 'API', 'Instagram Business account identifier'],
        ['username', 'STRING', 'API', 'Instagram handle'],
        ['account_id', 'STRING', 'API', 'Alternative account identifier'],
        ['page_id', 'STRING', 'API', 'Linked Facebook Page ID'],
        ['date', 'DATE', 'API', 'Calendar date of the metrics'],
        ['reach', 'FLOAT64', 'API', 'Unique accounts that saw any content from this account on this date'],
        ['follower_count', 'FLOAT64', 'API', 'Net new followers gained on this date (last 30 days only; older = 0)'],
        ['views', 'FLOAT64', 'Media aggregation', 'Total views on all posts published on this date (lifetime views summed from per-post data)'],
        ['likes', 'FLOAT64', 'Media aggregation', 'Total likes on posts published on this date'],
        ['comments', 'FLOAT64', 'Media aggregation', 'Total comments on posts published on this date'],
        ['shares', 'FLOAT64', 'Media aggregation', 'Total shares on posts published on this date'],
        ['saved', 'FLOAT64', 'Media aggregation', 'Total saves/bookmarks on posts published on this date'],
        ['total_interactions', 'FLOAT64', 'Media aggregation', 'Sum of all engagement actions on posts published on this date'],
        ['posts_published', 'INT64', 'Media aggregation', 'Number of posts published on this date. Use to distinguish "no engagement" from "no posts."'],
        ['engagement_rate', 'FLOAT64', 'Derived', 'total_interactions / reach. NULL if reach = 0. Measures how engaging the content was relative to audience size.'],
    ],
    col_widths=[3.5, 2, 3.5, 7.5]
)

doc.add_page_break()

# ============================================================
# 5. DATA REFRESH SCHEDULE
# ============================================================
add_heading('5. Data Refresh Schedule', level=1)
add_table(
    ['Component', 'Frequency', 'Method', 'Details'],
    [
        ['Account discovery', 'Daily', 'Full refresh', 'All 35 accounts re-extracted from Facebook Pages API'],
        ['Media + media insights', 'Daily', 'Incremental', '30-day lookback captures new posts and updated metrics for recent posts'],
        ['Account insights', 'Daily', 'Incremental', 'Reach: 30-day chunks. Follower count: last 30 days only (API limit)'],
        ['Stories + story insights', 'Daily', 'Full refresh', 'Only currently active stories available (24-hour window)'],
        ['Report tables', 'Daily', 'Full rebuild', 'CREATE OR REPLACE from source tables after each pipeline run'],
    ],
    col_widths=[4, 2.5, 3, 7]
)

# ============================================================
# 6. KNOWN LIMITATIONS
# ============================================================
add_heading('6. Known Limitations & API Constraints', level=1)

limitations = [
    ('Account-Level Views Derived from Media',
     'The views metric at the account level only supports aggregate responses from the Instagram API '
     '(total_value), not daily breakdowns. Rather than making one expensive API call per account per day '
     '(~25,550 calls for a 2-year backfill), daily views are derived by aggregating per-post views from '
     'media insights data we already collect. This gives "total views on posts published on date X" rather '
     'than "all views the account received on date X." Days with no posts show views = 0, clearly indicated '
     'by posts_published = 0.'),

    ('Follower Count: 30-Day Recency Limit',
     'The API only provides follower_count (net daily change) for the last 30 days. Historical follower '
     'data beyond 30 days cannot be retrieved. Once captured by our daily pipeline, the data is permanently '
     'preserved in BigQuery.'),

    ('Stories Are Ephemeral (24 Hours)',
     'Instagram Stories expire after 24 hours. The API only returns currently active stories. Story data '
     'depends on the pipeline running at least once per day. Historical stories that were not captured '
     'during their active window cannot be recovered.'),

    ('Deprecated Metrics (Meta 2025 Changes)',
     'Several metrics were removed by Meta:\n'
     '  - impressions: replaced by views (April 2025)\n'
     '  - profile_views, email_contacts, phone_call_clicks: removed entirely (January 2025)\n'
     '  - Story metrics navigation, total_interactions, saved, exits: removed (January 2025)'),

    ('Carousel Album Limitations',
     'CAROUSEL_ALBUM posts return 0 for likes, comments, and shares through the Insights API. '
     'This is a Meta API limitation, not a data quality issue. Native like_count and comments_count '
     'from the post object (instagram_media table) remain available for all media types including carousels.'),

    ('Data Type Convention',
     'All source table columns are stored as STRING type for schema consistency across the pipeline. '
     'Numeric fields are cast to FLOAT64 or INT64 in the report tables for computation. This design '
     'prevents schema mismatch errors during incremental loads.'),
]

for title, desc in limitations:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(8)

# ============================================================
# 7. UPCOMING CHANGES
# ============================================================
add_heading('7. Upcoming Changes', level=1)

p = doc.add_paragraph()
p.add_run('Reach to Viewers Migration (June 2026): ').bold = True
p.add_run(
    'Meta plans to replace the reach metric with viewers by June 2026. The new metric will measure '
    'unique people who actually viewed content (rather than had content delivered to their feed) and '
    'is expected to be 5-35% lower than current reach values. This will affect engagement_rate calculations. '
    'Our pipeline will be updated to use the new metric once Meta finalizes the transition.'
)

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Instagram_Data_Dictionary.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
